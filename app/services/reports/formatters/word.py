"""
Word Report Formatter

Generates Word documents with track changes simulation and native COM integration.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from ....utils.logging.setup import get_logger

logger = get_logger(__name__)

# Windows COM interface (optional)
try:
    import win32com.client as win32
    import pythoncom
    COM_AVAILABLE = True
except ImportError:
    COM_AVAILABLE = False
    logger.debug("Windows COM interface not available - Word track changes disabled")


class WordReportFormatter:
    """
    Word report formatter for contract analysis results.
    
    Supports:
    - Redlined documents with track changes simulation
    - Native Word track changes via COM (Windows only)
    - Professional document formatting
    """
    
    def __init__(self):
        """Initialize Word formatter"""
        self.com_available = COM_AVAILABLE
        if not self.com_available:
            logger.debug("Word COM not available - using docx library only")
    
    def generate_redlined_document(self, analysis_data: Dict[str, Any], output_path: str) -> str:
        """
        Generate Word redlined document with track changes simulation
        
        Args:
            analysis_data: Analysis results containing changes and metadata
            output_path: Path to save the Word document
            
        Returns:
            Path to the generated Word document
        """
        logger.info(f"Generating Word redlined document: {Path(output_path).name}")
        
        doc = Document()
        
        # Document title
        title = doc.add_heading('Contract Redlined Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive summary section
        doc.add_heading('Executive Summary', level=1)
        self._add_executive_summary(doc, analysis_data)
        
        # Document comparison section
        doc.add_heading('Document Changes', level=1)
        self._add_changes_section(doc, analysis_data)
        
        # Risk assessment section
        doc.add_heading('Risk Assessment', level=1)
        self._add_risk_assessment(doc, analysis_data)
        
        # Recommendations section
        doc.add_heading('Recommendations', level=1)
        self._add_recommendations(doc, analysis_data)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Word document saved: {output_path}")
        
        return output_path
    
    def generate_word_com_redlined(self, analysis_data: Dict[str, Any], output_path: str) -> Optional[str]:
        """
        Generate Word document with native track changes using COM (Windows only)
        
        Args:
            analysis_data: Analysis results
            output_path: Path to save the Word document
            
        Returns:
            Path to generated file or None if COM not available
        """
        if not self.com_available:
            logger.warning("Windows COM interface not available - cannot generate native track changes")
            return None
        
        # Get contract and template paths
        contract_path = analysis_data.get('contract_path')
        template_path = analysis_data.get('template_path')
        
        if not contract_path or not template_path:
            logger.error("Contract or template path missing from analysis data")
            return None
        
        if not Path(contract_path).exists() or not Path(template_path).exists():
            logger.error("Contract or template file not found")
            return None
        
        return self._generate_com_redlined_document(
            template_path, contract_path, output_path, analysis_data
        )
    
    def _generate_com_redlined_document(
        self, 
        template_path: str, 
        contract_path: str, 
        output_path: str,
        analysis_data: Dict[str, Any]
    ) -> Optional[str]:
        """Generate redlined document using Word COM interface"""
        
        # COM objects to clean up
        word = None
        original = None
        revised = None
        compared_doc = None
        
        try:
            # Initialize COM
            pythoncom.CoInitialize()
            
            logger.info(f"Starting Word COM comparison: {Path(template_path).name} vs {Path(contract_path).name}")
            
            # Launch Word
            word = win32.Dispatch('Word.Application')
            word.Visible = False
            
            # Open documents
            original = word.Documents.Open(str(Path(template_path).resolve()))
            revised = word.Documents.Open(str(Path(contract_path).resolve()))
            
            # Compare documents to create redlined version
            compared_doc = word.CompareDocuments(
                OriginalDocument=original,
                RevisedDocument=revised,
                Destination=win32.constants.wdCompareDestinationNew,
                Granularity=win32.constants.wdGranularityWordLevel,
                CompareFormatting=True,
                CompareCaseChanges=True,
                CompareWhitespace=True,
                CompareTables=True,
                CompareHeaders=True,
                CompareFootnotes=True,
                CompareTextboxes=True,
                CompareFields=True,
                CompareComments=True,
                CompareMoves=True,
                RevisedAuthor="Contract Analyzer",
                OriginalAuthor="Template"
            )
            
            # Ensure track changes are visible
            compared_doc.TrackRevisions = True
            compared_doc.ShowRevisions = True
            
            # Add header with analysis info
            if compared_doc.Sections.Count > 0:
                section = compared_doc.Sections(1)
                if section.Headers.Count > 0:
                    header = section.Headers(win32.constants.wdHeaderFooterPrimary)
                    header_range = header.Range
                    header_range.Text = f"Contract Analysis Report - Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
                    header_range.Font.Size = 9
                    header_range.Font.Name = "Arial"
            
            # Save the compared document
            compared_doc.SaveAs2(str(Path(output_path).resolve()))
            
            logger.info(f"Word COM redlined document generated: {Path(output_path).name}")
            return output_path
            
        except Exception as e:
            logger.error(f"Word COM redlined document generation failed: {e}")
            return None
            
        finally:
            # Clean up COM objects
            self._safe_com_cleanup(word, original, revised, compared_doc)
            
            # Always cleanup COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass
    
    def _safe_com_cleanup(self, word=None, original=None, revised=None, compared_doc=None):
        """Safely clean up COM objects without propagating errors"""
        
        # Close documents in reverse order (newest first)
        documents_to_close = [
            ("compared_doc", compared_doc),
            ("revised", revised), 
            ("original", original)
        ]
        
        for doc_name, doc in documents_to_close:
            if doc is not None:
                try:
                    logger.debug(f"Closing {doc_name}...")
                    doc.Close(False)  # Don't save changes
                except Exception as e:
                    logger.debug(f"Warning: Could not close {doc_name}: {e}")
        
        # Quit Word application
        if word is not None:
            try:
                logger.debug("Quitting Word application...")
                word.Quit()
            except Exception as e:
                logger.debug(f"Warning: Could not quit Word cleanly: {e}")
        
        logger.debug("COM cleanup completed")
    
    def _add_executive_summary(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add executive summary section to the document"""
        
        summary_para = doc.add_paragraph()
        summary_para.add_run("Analysis Date: ").bold = True
        summary_para.add_run(f"{analysis_data.get('date', datetime.now().strftime('%Y-%m-%d'))}\n")
        
        summary_para.add_run("Contract: ").bold = True
        summary_para.add_run(f"{analysis_data.get('contract', 'Unknown')}\n")
        
        summary_para.add_run("Template: ").bold = True
        summary_para.add_run(f"{analysis_data.get('template', 'Unknown')}\n")
        
        summary_para.add_run("Total Changes: ").bold = True
        summary_para.add_run(f"{analysis_data.get('changes', 0)}\n")
        
        summary_para.add_run("Similarity Score: ").bold = True
        summary_para.add_run(f"{analysis_data.get('similarity', 0)}%\n")
        
        # Risk level with color coding
        risk_level = self._determine_risk_level(analysis_data)
        summary_para.add_run("Risk Level: ").bold = True
        risk_run = summary_para.add_run(f"{risk_level}\n")
        
        # Color code risk level
        if risk_level == "HIGH":
            risk_run.font.color.rgb = RGBColor(204, 0, 0)  # Red
        elif risk_level == "MEDIUM":
            risk_run.font.color.rgb = RGBColor(255, 102, 0)  # Orange
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        risk_run.bold = True
    
    def _add_changes_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add changes section with track changes simulation"""
        
        changes = analysis_data.get('analysis', [])
        if not changes:
            doc.add_paragraph("No changes detected in this document.")
            return
        
        # Group changes by criticality
        changes_by_type = self._group_changes_by_type(changes)
        
        for criticality in ['critical', 'significant', 'inconsequential']:
            type_changes = changes_by_type.get(criticality, [])
            if not type_changes:
                continue
            
            # Section heading
            doc.add_heading(f"{criticality.title()} Changes ({len(type_changes)})", level=2)
            
            for i, change in enumerate(type_changes, 1):
                change_para = doc.add_paragraph()
                
                # Change number
                change_para.add_run(f"{i}. ").bold = True
                
                # Deleted text (strikethrough simulation)
                deleted_text = change.get('deleted_text', '').strip()
                if deleted_text:
                    deleted_run = change_para.add_run(f"[DELETED: {deleted_text}] ")
                    deleted_run.font.color.rgb = RGBColor(204, 0, 0)  # Red
                    deleted_run.font.strike = True
                
                # Inserted text (underlined)
                inserted_text = change.get('inserted_text', '').strip()
                if inserted_text:
                    inserted_run = change_para.add_run(f"[INSERTED: {inserted_text}] ")
                    inserted_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    inserted_run.underline = True
                
                # Explanation
                explanation = change.get('explanation', 'No explanation provided')
                change_para.add_run(f"\nExplanation: {explanation}\n")
    
    def _add_risk_assessment(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add risk assessment section"""
        
        risk_level = self._determine_risk_level(analysis_data)
        risk_explanation = self._get_risk_explanation(risk_level, analysis_data)
        
        risk_para = doc.add_paragraph()
        risk_para.add_run("Overall Risk Level: ").bold = True
        
        risk_run = risk_para.add_run(f"{risk_level}\n\n")
        risk_run.bold = True
        
        # Color code risk level
        if risk_level == "HIGH":
            risk_run.font.color.rgb = RGBColor(204, 0, 0)  # Red
        elif risk_level == "MEDIUM":
            risk_run.font.color.rgb = RGBColor(255, 102, 0)  # Orange
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        risk_para.add_run(risk_explanation)
    
    def _add_recommendations(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add recommendations section"""
        
        recommendations = self._generate_recommendations(analysis_data)
        
        if recommendations:
            for rec in recommendations:
                rec_para = doc.add_paragraph(rec, style='List Bullet')
        else:
            doc.add_paragraph("No specific recommendations at this time.")
    
    def _group_changes_by_type(self, changes: List[Dict]) -> Dict[str, List[Dict]]:
        """Group changes by criticality type"""
        groups = {
            'critical': [],
            'significant': [],
            'inconsequential': []
        }
        
        for change in changes:
            classification = change.get('classification', 'INCONSEQUENTIAL').lower()
            if classification in groups:
                groups[classification].append(change)
            else:
                groups['inconsequential'].append(change)
        
        return groups
    
    def _determine_risk_level(self, analysis_data: Dict[str, Any]) -> str:
        """Determine overall risk level"""
        changes = analysis_data.get('analysis', [])
        critical_count = len([c for c in changes if c.get('classification') == 'CRITICAL'])
        significant_count = len([c for c in changes if c.get('classification') == 'SIGNIFICANT'])
        
        if critical_count > 0:
            return "HIGH"
        elif significant_count > 5:
            return "HIGH"
        elif significant_count > 0:
            return "MEDIUM"
        else:
            return "LOW"
    
    def _get_risk_explanation(self, risk_level: str, analysis_data: Dict[str, Any]) -> str:
        """Get explanation for risk level"""
        changes = analysis_data.get('analysis', [])
        critical_count = len([c for c in changes if c.get('classification') == 'CRITICAL'])
        
        if risk_level == "HIGH":
            if critical_count > 0:
                return f"This contract contains {critical_count} critical change(s) involving actual value-to-value modifications (e.g., price changes, service changes, company changes). These require immediate legal review and approval before proceeding."
            else:
                return "This contract contains numerous significant changes that require immediate legal review and approval before proceeding."
        elif risk_level == "MEDIUM":
            return "This contract contains significant changes that should be reviewed by legal counsel before execution."
        else:
            return "This contract has minimal changes (mostly placeholder → actual content) and can proceed through standard review processes."
    
    def _generate_recommendations(self, analysis_data: Dict[str, Any]) -> List[str]:
        """Generate recommendations based on analysis"""
        recommendations = []
        
        changes = analysis_data.get('analysis', [])
        critical_changes = [c for c in changes if c.get('classification') == 'CRITICAL']
        risk_level = self._determine_risk_level(analysis_data)
        
        if risk_level == "HIGH":
            recommendations.extend([
                "🚨 CRITICAL: Schedule immediate legal review with qualified counsel",
                "🚨 Do not execute contract until all critical changes are approved",
                "🚨 Focus on value-to-value changes (price, service, company modifications)"
            ])
        elif risk_level == "MEDIUM":
            recommendations.extend([
                "⚠️ Schedule legal review before contract execution",
                "⚠️ Review all significant changes with stakeholders",
                "⚠️ Verify pricing and service level changes"
            ])
        else:
            recommendations.extend([
                "✅ Contract may proceed through standard review process",
                "✅ Verify placeholder content has been properly filled",
                "✅ Confirm standard terms and conditions"
            ])
        
        # Add general recommendations
        recommendations.extend([
            "📋 Document all approved changes for future reference",
            "📋 Ensure all parties acknowledge the modifications",
            "📋 Update contract management system with new terms"
        ])
        
        return recommendations


__all__ = ['WordReportFormatter']