"""
Document Processor Service

Handles document text extraction and processing operations.
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document

from ...utils.logging.setup import get_logger

logger = get_logger(__name__)


class DocumentProcessingError(Exception):
    """Exception raised when document processing fails"""
    pass


class DocumentProcessor:
    """
    Document processing service for extracting and manipulating document content.
    
    Handles:
    - Text extraction from DOCX files
    - Document structure analysis
    - Metadata extraction
    - Text preprocessing and cleaning
    """
    
    def __init__(self):
        """Initialize document processor"""
        logger.debug("Document processor initialized")
    
    def extract_text_from_docx(self, filepath: str) -> str:
        """
        Extract text content from a .docx file.
        
        Args:
            filepath: Path to the .docx file
            
        Returns:
            Full text content as a single string
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            if not os.path.exists(filepath):
                raise DocumentProcessingError(f"File not found: {filepath}")
            
            doc = Document(filepath)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
            logger.debug(f"Extracted {len(full_text)} characters from {Path(filepath).name}")
            return full_text
            
        except Exception as e:
            error_msg = f"Error extracting text from {filepath}: {str(e)}"
            logger.error(error_msg)
            raise DocumentProcessingError(error_msg)
    
    def extract_structured_content(self, filepath: str) -> Dict[str, Any]:
        """
        Extract structured content from a DOCX file including metadata
        
        Args:
            filepath: Path to the .docx file
            
        Returns:
            Dictionary with structured content
        """
        try:
            doc = Document(filepath)
            
            content = {
                'paragraphs': [],
                'tables': [],
                'headings': [],
                'metadata': {
                    'paragraph_count': 0,
                    'table_count': 0,
                    'heading_count': 0,
                    'word_count': 0,
                    'character_count': 0
                }
            }
            
            # Extract paragraphs and identify headings
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    para_data = {
                        'text': para_text,
                        'style': para.style.name if para.style else 'Normal'
                    }
                    
                    # Check if it's a heading
                    if para.style and 'Heading' in para.style.name:
                        content['headings'].append(para_data)
                        content['metadata']['heading_count'] += 1
                    
                    content['paragraphs'].append(para_data)
                    content['metadata']['paragraph_count'] += 1
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                content['tables'].append(table_data)
                content['metadata']['table_count'] += 1
            
            # Calculate text statistics
            full_text = self.extract_text_from_docx(filepath)
            content['metadata']['character_count'] = len(full_text)
            content['metadata']['word_count'] = len(full_text.split())
            
            logger.debug(f"Extracted structured content from {Path(filepath).name}")
            return content
            
        except Exception as e:
            error_msg = f"Error extracting structured content from {filepath}: {str(e)}"
            logger.error(error_msg)
            raise DocumentProcessingError(error_msg)
    
    def preprocess_text(self, text: str) -> str:
        """
        Preprocess text for analysis
        
        Args:
            text: Raw text to preprocess
            
        Returns:
            Preprocessed text
        """
        if not text:
            return ""
        
        # Basic text cleaning
        processed_text = text.strip()
        
        # Normalize whitespace
        processed_text = ' '.join(processed_text.split())
        
        # Remove excessive line breaks
        processed_text = processed_text.replace('\n\n\n', '\n\n')
        
        return processed_text
    
    def create_commented_docx(
        self,
        original_filepath: str,
        analysis_results: List[Dict[str, Any]],
        output_path: str
    ) -> bool:
        """
        Create a commented version of the original document with AI analysis.
        
        Args:
            original_filepath: Path to the original uploaded document
            analysis_results: List of analysis results from LLM
            output_path: Path for the output commented document
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Load the original document
            doc = Document(original_filepath)
            
            # Add a new section for analysis results
            doc.add_heading('Contract Analysis Summary', level=1)
            
            # Categorize changes
            critical_changes = [r for r in analysis_results if r.get('classification') == 'CRITICAL']
            significant_changes = [r for r in analysis_results if r.get('classification') == 'SIGNIFICANT']
            inconsequential_changes = [r for r in analysis_results if r.get('classification') == 'INCONSEQUENTIAL']
            
            # Add summary paragraph
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Analysis completed on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            summary_para.add_run(f"Total changes detected: {len(analysis_results)}\n")
            summary_para.add_run(f"Critical changes: {len(critical_changes)}\n")
            summary_para.add_run(f"Significant changes: {len(significant_changes)}\n")
            summary_para.add_run(f"Inconsequential changes: {len(inconsequential_changes)}\n")
            
            # Add critical changes section
            if critical_changes:
                doc.add_heading('Critical Changes', level=2)
                for i, change in enumerate(critical_changes, 1):
                    self._add_change_to_document(doc, i, change)
            
            # Add significant changes section
            if significant_changes:
                doc.add_heading('Significant Changes', level=2)
                for i, change in enumerate(significant_changes, 1):
                    self._add_change_to_document(doc, i, change)
            
            # Add inconsequential changes section
            if inconsequential_changes:
                doc.add_heading('Inconsequential Changes', level=2)
                for i, change in enumerate(inconsequential_changes, 1):
                    self._add_change_to_document(doc, i, change)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Commented document created: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating commented document: {str(e)}")
            return False
    
    def _add_change_to_document(self, doc: Document, index: int, change: Dict[str, Any]):
        """Add a single change to the document"""
        para = doc.add_paragraph()
        para.add_run(f"{index}. ").bold = True
        para.add_run(change.get('explanation', 'No explanation provided'))
        
        # Add deleted text
        if change.get('deleted_text'):
            para = doc.add_paragraph()
            para.add_run("Deleted: ").bold = True
            deleted_text = change['deleted_text']
            if len(deleted_text) > 200:
                deleted_text = deleted_text[:200] + "..."
            para.add_run(deleted_text)
        
        # Add inserted text
        if change.get('inserted_text'):
            para = doc.add_paragraph()
            para.add_run("Added: ").bold = True
            inserted_text = change['inserted_text']
            if len(inserted_text) > 200:
                inserted_text = inserted_text[:200] + "..."
            para.add_run(inserted_text)
        
        doc.add_paragraph()  # Add spacing
    
    def save_analysis_metadata(
        self,
        filename: str,
        template_used: str,
        analysis_results: List[Dict[str, Any]],
        output_dir: str
    ) -> str:
        """
        Save analysis metadata to JSON file.
        
        Args:
            filename: Original filename
            template_used: Template that was used for comparison
            analysis_results: Analysis results from LLM
            output_dir: Directory to save metadata
            
        Returns:
            Path to saved metadata file
        """
        try:
            metadata = {
                'uploaded_filename': filename,
                'selected_template': template_used,
                'analysis_timestamp': datetime.now().isoformat(),
                'total_changes': len(analysis_results),
                'critical_changes_count': len([r for r in analysis_results if r.get('classification') == 'CRITICAL']),
                'significant_changes_count': len([r for r in analysis_results if r.get('classification') == 'SIGNIFICANT']),
                'inconsequential_changes_count': len([r for r in analysis_results if r.get('classification') == 'INCONSEQUENTIAL']),
                'analysis_results': analysis_results
            }
            
            # Create metadata filename
            base_name = os.path.splitext(filename)[0]
            metadata_path = os.path.join(output_dir, f"{base_name}_analysis.json")
            
            # Save metadata
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            logger.info(f"Analysis metadata saved: {metadata_path}")
            return metadata_path
            
        except Exception as e:
            logger.error(f"Error saving analysis metadata: {str(e)}")
            raise DocumentProcessingError(f"Failed to save metadata: {e}")
    
    def validate_document(self, filepath: str) -> Dict[str, Any]:
        """
        Validate a document file
        
        Args:
            filepath: Path to document file
            
        Returns:
            Validation result with status and details
        """
        result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'metadata': {}
        }
        
        try:
            path = Path(filepath)
            
            # Check file existence
            if not path.exists():
                result['errors'].append("File does not exist")
                return result
            
            # Check file extension
            if path.suffix.lower() not in ['.docx', '.doc']:
                result['errors'].append("Invalid file type. Only .docx and .doc files are supported")
                return result
            
            # Check file size
            file_size = path.stat().st_size
            max_size = 50 * 1024 * 1024  # 50MB limit
            if file_size > max_size:
                result['errors'].append(f"File too large. Maximum size is {max_size // (1024*1024)}MB")
                return result
            
            # Try to open and read document
            try:
                text_content = self.extract_text_from_docx(filepath)
                if not text_content.strip():
                    result['warnings'].append("Document appears to be empty")
                
                result['metadata'] = {
                    'file_size': file_size,
                    'character_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'has_content': bool(text_content.strip())
                }
                
            except Exception as e:
                result['errors'].append(f"Unable to read document content: {str(e)}")
                return result
            
            # If we get here, document is valid
            result['valid'] = True
            
        except Exception as e:
            result['errors'].append(f"Document validation failed: {str(e)}")
        
        return result


__all__ = ['DocumentProcessor', 'DocumentProcessingError']