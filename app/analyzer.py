# app/analyzer.py
import difflib
from docx import Document
import os
import json
from datetime import datetime

class ContractAnalyzer:
    """
    Contract Analyzer class that provides document analysis functionality.
    """
    
    def __init__(self):
        """Initialize the ContractAnalyzer."""
        pass
    
    def extract_text_from_docx(self, filepath):
        """
        Extract text content from a .docx file.
        
        Args:
            filepath (str): Path to the .docx file
            
        Returns:
            str: Full text content as a single string
        """
        return extract_text_from_docx(filepath)
    
    def find_changes(self, template_text, contract_text):
        """
        Find changes between template and contract text.
        
        Args:
            template_text (str): Original template text
            contract_text (str): Contract text to compare
            
        Returns:
            list: List of changes
        """
        return compare_texts(template_text, contract_text)
    
    def calculate_similarity(self, template_text, contract_text):
        """
        Calculate similarity between template and contract text.
        
        Args:
            template_text (str): Original template text
            contract_text (str): Contract text to compare
            
        Returns:
            float: Similarity ratio (0.0 to 1.0)
        """
        return difflib.SequenceMatcher(None, template_text, contract_text).ratio()
    
    def create_commented_docx(self, original_filepath, analysis_results, output_path):
        """
        Create a commented version of the original document with AI analysis.
        
        Args:
            original_filepath (str): Path to the original uploaded document
            analysis_results (list): List of analysis results from LLM
            output_path (str): Path for the output commented document
            
        Returns:
            bool: True if successful, False otherwise
        """
        return create_commented_docx(original_filepath, analysis_results, output_path)
    
    def save_analysis_metadata(self, filename, template_used, analysis_results, output_dir):
        """
        Save analysis metadata to JSON file.
        
        Args:
            filename (str): Original filename
            template_used (str): Template that was used for comparison
            analysis_results (list): Analysis results from LLM
            output_dir (str): Directory to save metadata
            
        Returns:
            str: Path to saved metadata file
        """
        return save_analysis_metadata(filename, template_used, analysis_results, output_dir)

def extract_text_from_docx(filepath):
    """
    Extract text content from a .docx file.
    
    Args:
        filepath (str): Path to the .docx file
        
    Returns:
        str: Full text content as a single string
    """
    try:
        doc = Document(filepath)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
            
        return '\n'.join(text_content)
    except Exception as e:
        print(f"Error extracting text from {filepath}: {str(e)}")
        return ""

def compare_texts(text1, text2):
    """
    Compare two texts and return structured differences.
    
    Args:
        text1 (str): Original text
        text2 (str): Modified text
        
    Returns:
        list: List of differences in format [('operation', 'text'), ...]
    """
    differ = difflib.unified_diff(
        text1.splitlines(keepends=True),
        text2.splitlines(keepends=True),
        lineterm='',
        n=0
    )
    
    changes = []
    for line in differ:
        if line.startswith('--- ') or line.startswith('+++ ') or line.startswith('@@'):
            continue
        elif line.startswith('-'):
            changes.append(('delete', line[1:]))
        elif line.startswith('+'):
            changes.append(('insert', line[1:]))
    
    return changes

def create_commented_docx(original_filepath, analysis_results, output_path):
    """
    Create a commented version of the original document with AI analysis.
    
    Args:
        original_filepath (str): Path to the original uploaded document
        analysis_results (list): List of analysis results from LLM
        output_path (str): Path for the output commented document
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Load the original document
        doc = Document(original_filepath)
        
        # For now, we'll add comments as a summary at the end
        # In a more sophisticated version, we'd match text and add inline comments
        
        # Add a new section for analysis results
        doc.add_heading('Contract Analysis Summary', level=1)
        
        significant_changes = [r for r in analysis_results if r.get('classification') == 'SIGNIFICANT']
        inconsequential_changes = [r for r in analysis_results if r.get('classification') == 'INCONSEQUENTIAL']
        
        # Add summary paragraph
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Analysis completed on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        summary_para.add_run(f"Total changes detected: {len(analysis_results)}\n")
        summary_para.add_run(f"Significant changes: {len(significant_changes)}\n")
        summary_para.add_run(f"Inconsequential changes: {len(inconsequential_changes)}\n")
        
        # Add significant changes section
        if significant_changes:
            doc.add_heading('Significant Changes', level=2)
            for i, change in enumerate(significant_changes, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(change.get('explanation', 'No explanation provided'))
                
                if change.get('deleted_text'):
                    para = doc.add_paragraph()
                    para.add_run("Deleted: ").bold = True
                    para.add_run(change['deleted_text'][:200] + "..." if len(change['deleted_text']) > 200 else change['deleted_text'])
                
                if change.get('inserted_text'):
                    para = doc.add_paragraph()
                    para.add_run("Added: ").bold = True
                    para.add_run(change['inserted_text'][:200] + "..." if len(change['inserted_text']) > 200 else change['inserted_text'])
                
                doc.add_paragraph()  # Add spacing
        
        # Add inconsequential changes section
        if inconsequential_changes:
            doc.add_heading('Inconsequential Changes', level=2)
            for i, change in enumerate(inconsequential_changes, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(change.get('explanation', 'No explanation provided'))
        
        # Save the document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating commented document: {str(e)}")
        return False

def save_analysis_metadata(filename, template_used, analysis_results, output_dir):
    """
    Save analysis metadata to JSON file.
    
    Args:
        filename (str): Original filename
        template_used (str): Template that was used for comparison
        analysis_results (list): Analysis results from LLM
        output_dir (str): Directory to save metadata
        
    Returns:
        str: Path to saved metadata file
    """
    metadata = {
        'uploaded_filename': filename,
        'selected_template': template_used,
        'analysis_timestamp': datetime.now().isoformat(),
        'total_changes': len(analysis_results),
        'significant_changes_count': len([r for r in analysis_results if r.get('classification') == 'SIGNIFICANT']),
        'analysis_results': analysis_results
    }
    
    # Create metadata filename
    base_name = os.path.splitext(filename)[0]
    metadata_path = os.path.join(output_dir, f"{base_name}_analysis.json")
    
    # Save metadata
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=2)
    
    return metadata_path 