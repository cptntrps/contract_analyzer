#!/usr/bin/env python3
"""
Advanced Contract Analyzer
A tool for automated contract review with semantic template matching and comprehensive reporting.
"""

import os
import sys
import json
import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import concurrent.futures
from dataclasses import dataclass

import docx
import difflib
from sentence_transformers import SentenceTransformer, util
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import torch


@dataclass
class AnalysisResult:
    """Data class to store contract analysis results"""
    input_file: str
    matched_template: str
    similarity_score: float
    category: str
    differences_count: int
    report_path: str
    timestamp: str


class ContractAnalyzer:
    """Main contract analyzer class with batch processing capabilities"""
    
    def __init__(self, template_dir: str = 'templates', model_name: str = 'all-MiniLM-L6-v2', use_ml: bool = False):
        self.template_dir = Path(template_dir)
        self.use_ml = use_ml
        self.templates = {}
        
        # Only initialize ML model if explicitly requested
        if self.use_ml:
            self.model = SentenceTransformer(model_name)
            self.template_embeddings = {}
        else:
            self.model = None
            self.template_embeddings = None
        
        # Enhanced keyword definitions with regex support
        self.vendor_keywords = {
            "capgemini": "VENDOR_CAPGEMINI",
            "blue optima": "VENDOR_BLUEOPTIMA", 
            "blueoptima": "VENDOR_BLUEOPTIMA",
            "epam": "VENDOR_EPAM",
            "accenture": "VENDOR_ACCENTURE",
            "ibm": "VENDOR_IBM",
            "tcs": "VENDOR_TCS",
            "infosys": "VENDOR_INFOSYS",
            "wipro": "VENDOR_WIPRO",
        }
        
        # Prioritized keyword list - order matters for better classification
        self.type_keywords = {
            # High-priority specific identifiers first
            "change order template": "TYPE_CHANGEORDER",
            "change order number": "TYPE_CHANGEORDER", 
            "statement of work template": "TYPE_SOW",
            "change order": "TYPE_CHANGEORDER",
            "change request": "TYPE_CHANGEORDER",
            "master service agreement": "TYPE_MSA",
            "non-disclosure agreement": "TYPE_NDA",
            "service level agreement": "TYPE_SLA",
            "purchase order": "TYPE_PO",
            # Lower priority - often just references
            "statement of work": "TYPE_SOW",  
            "msa": "TYPE_MSA",
            "nda": "TYPE_NDA",
            "sla": "TYPE_SLA",
            "sow": "TYPE_SOW",  # Lowest priority - very common in references
        }
        
        self.load_templates()
    
    def load_templates(self):
        """Load and process all templates from the template directory"""
        if not self.template_dir.exists():
            print(f"Warning: Template directory '{self.template_dir}' does not exist. Creating it...")
            self.template_dir.mkdir(parents=True)
            return
        
        template_files = list(self.template_dir.glob("*.docx"))
        if not template_files:
            print(f"Warning: No .docx templates found in '{self.template_dir}'")
            return
        
        print(f"Loading {len(template_files)} templates...")
        for template_path in template_files:
            text = self.extract_text_from_docx(str(template_path))
            if text:
                self.templates[template_path.name] = text
                print(f"  - Loaded: {template_path.name}")
                # Pre-compute embeddings only if using ML
                if self.use_ml and self.model:
                    self.template_embeddings[template_path.name] = self.model.encode(text)
        
        print(f"Successfully loaded {len(self.templates)} templates")
    
    def extract_text_from_docx(self, filepath: str, skip_initial_table: bool = True) -> str:
        """
        Extract text from a .docx file with improved error handling
        """
        try:
            doc = docx.Document(filepath)
            
            # Handle empty documents
            if not doc.paragraphs and not doc.tables:
                return ""
            
            # If no tables or skip_initial_table is False, return all text
            if not doc.tables or not skip_initial_table:
                return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            
            # Skip initial table logic
            body_paragraphs = []
            start_collecting = False
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = docx.text.paragraph.Paragraph(element, doc)
                    if start_collecting and para.text.strip():
                        body_paragraphs.append(para.text.strip())
                elif element.tag.endswith('tbl'):  # Table
                    start_collecting = True
            
            return "\n".join(body_paragraphs)
            
        except Exception as e:
            print(f"Error reading {filepath}: {e}")
            return ""
    
    def classify_document(self, text: str) -> Optional[str]:
        """Enhanced document classification with frequency scoring"""
        if not text:
            return None
            
        lower_text = text.lower()
        
        # Check vendor keywords first (highest priority)
        for keyword, category in self.vendor_keywords.items():
            if keyword in lower_text:
                print(f"  Vendor classification: {category} (keyword: '{keyword}')")
                return category
        
        # Count keyword frequencies for each category
        category_scores = {}
        keyword_details = []
        
        for keyword, category in self.type_keywords.items():
            count = lower_text.count(keyword)
            if count > 0:
                # Give higher weight to more specific keywords
                weight = len(keyword.split())  # Multi-word keywords get higher weight
                if "template" in keyword:
                    weight *= 5  # "template" indicates primary document type
                elif keyword in ["change order", "change request"]:
                    weight *= 3  # Specific document types get higher weight
                elif keyword in ["sow", "msa", "nda", "sla"]:
                    weight *= 0.5  # Abbreviations often used as references
                
                score = count * weight
                
                if category not in category_scores:
                    category_scores[category] = 0
                category_scores[category] += score
                
                keyword_details.append(f"'{keyword}': {count} × {weight:.1f} = {score:.1f}")
        
        if category_scores:
            # Find the category with highest score
            best_category = max(category_scores.items(), key=lambda x: x[1])[0]
            best_score = category_scores[best_category]
            
            print(f"  Classification analysis:")
            for detail in keyword_details:
                print(f"    {detail}")
            print(f"  Category scores: {category_scores}")
            print(f"  Document type: {best_category} (score: {best_score:.1f})")
            
            return best_category
        
        print("  No document type classification found")
        return None
    
    def find_best_template(self, input_text: str, category: str) -> Tuple[str, float]:
        """Find the best matching template using exact text similarity (like Word)"""
        # Filter templates by category
        relevant_templates = {
            name: text for name, text in self.templates.items()
            if name.upper().startswith(category)
        }
        
        if not relevant_templates:
            return None, 0.0
        
        # Use simple text similarity like Word track changes
        similarities = {}
        
        for template_name, template_text in relevant_templates.items():
            # Calculate text similarity ratio (like Word does)
            matcher = difflib.SequenceMatcher(None, template_text, input_text)
            similarity = matcher.ratio()
            similarities[template_name] = similarity
        
        # Get best match
        best_template = max(similarities.items(), key=lambda x: x[1])
        return best_template
    
    def generate_comparison_report(self, base_text: str, new_text: str, 
                                 output_path: str, contract_name: str) -> int:
        """Generate Word-style track changes comparison report"""
        doc = docx.Document()
        
        # Add header
        header = doc.add_heading('Contract Comparison Report', 0)
        header.alignment = 1  # Center
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Contract: ").bold = True
        metadata_para.add_run(f"{contract_name}\n")
        metadata_para.add_run(f"Generated: ").bold = True
        metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        metadata_para.add_run(f"Template Comparison: ").bold = True
        metadata_para.add_run("Approved Template vs. Submitted Contract\n")
        doc.add_paragraph()
        
        # Calculate differences first
        base_lines = base_text.splitlines()
        new_lines = new_text.splitlines()
        matcher = difflib.SequenceMatcher(None, base_lines, new_lines)
        
        added_count = 0
        deleted_count = 0
        modified_count = 0
        total_changes = 0
        
        # Count changes for summary
        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == 'insert':
                added_count += j2 - j1
                total_changes += j2 - j1
            elif opcode == 'delete':
                deleted_count += i2 - i1
                total_changes += i2 - i1
            elif opcode == 'replace':
                modified_count += max(i2 - i1, j2 - j1)
                total_changes += max(i2 - i1, j2 - j1)
        
        # Add summary section with track changes legend
        doc.add_heading('Change Summary & Legend', level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Changes Detected: {total_changes}\n").bold = True
        
        # Legend with actual formatting examples
        legend_para = doc.add_paragraph()
        legend_para.add_run("Legend:\n").bold = True
        
        # Deleted text example
        deleted_run = legend_para.add_run("■ Deleted Text")
        deleted_run.font.strike = True
        deleted_run.font.color.rgb = RGBColor(255, 0, 0)
        legend_para.add_run(f" - Text removed from template ({deleted_count} deletions)\n")
        
        # Inserted text example  
        inserted_run = legend_para.add_run("■ Inserted Text")
        inserted_run.font.color.rgb = RGBColor(0, 128, 0)
        inserted_run.font.underline = True
        legend_para.add_run(f" - Text added to contract ({added_count} additions)\n")
        
        # Modified text example
        modified_run = legend_para.add_run("■ Modified Text")
        modified_run.font.color.rgb = RGBColor(255, 165, 0)
        modified_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        legend_para.add_run(f" - Text changed from template ({modified_count} modifications)\n")
        
        doc.add_paragraph()
        
        # Add detailed comparison with improved formatting
        doc.add_heading('Detailed Track Changes View', level=1)
        
        # Create sentence-level comparison for better readability
        base_sentences = self._split_into_sentences(base_text)
        new_sentences = self._split_into_sentences(new_text)
        sentence_matcher = difflib.SequenceMatcher(None, base_sentences, new_sentences)
        
        comparison_para = doc.add_paragraph()
        
        for opcode, i1, i2, j1, j2 in sentence_matcher.get_opcodes():
            if opcode == 'equal':
                # Unchanged text - normal formatting
                text = ' '.join(base_sentences[i1:i2])
                if text.strip():
                    comparison_para.add_run(text + ' ')
                    
            elif opcode == 'delete':
                # Deleted text - red strikethrough
                text = ' '.join(base_sentences[i1:i2])
                if text.strip():
                    run = comparison_para.add_run(text + ' ')
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    
            elif opcode == 'insert':
                # Inserted text - green underline
                text = ' '.join(new_sentences[j1:j2])
                if text.strip():
                    run = comparison_para.add_run(text + ' ')
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    run.font.underline = True
                    
            elif opcode == 'replace':
                # Modified text - show both versions
                base_text_part = ' '.join(base_sentences[i1:i2])
                new_text_part = ' '.join(new_sentences[j1:j2])
                
                if base_text_part.strip():
                    # Show deleted version
                    run = comparison_para.add_run(base_text_part + ' ')
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                
                if new_text_part.strip():
                    # Show inserted version
                    run = comparison_para.add_run(new_text_part + ' ')
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    run.font.underline = True
        
        # Add section break for better readability
        doc.add_paragraph()
        doc.add_heading('Change Analysis', level=1)
        
        analysis_para = doc.add_paragraph()
        if total_changes == 0:
            run = analysis_para.add_run("✓ COMPLIANT: No unauthorized changes detected.")
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.bold = True
        elif total_changes <= 10:
            run = analysis_para.add_run("⚠ MINOR CHANGES: Few modifications detected - review recommended.")
            run.font.color.rgb = RGBColor(255, 165, 0)
            run.bold = True
        else:
            run = analysis_para.add_run("❌ MAJOR CHANGES: Significant modifications detected - thorough review required.")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
        
        # Save document
        doc.save(output_path)
        return total_changes
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better diff granularity"""
        import re
        # Simple sentence splitting - can be enhanced with nltk if needed
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def analyze_single_contract(self, input_path: str, output_dir: str = "reports") -> AnalysisResult:
        """Analyze a single contract file"""
        input_file = Path(input_path)
        if not input_file.exists():
            print(f"Error: File '{input_path}' not found")
            return None
        
        print(f"\nAnalyzing: {input_file.name}")
        
        # Extract text
        input_text = self.extract_text_from_docx(str(input_file))
        if not input_text:
            print(f"Error: Could not extract text from {input_file.name}")
            return None
        
        # Classify document
        category = self.classify_document(input_text)
        if not category:
            print(f"Warning: Could not classify {input_file.name}")
            return None
        
        print(f"  Category: {category}")
        
        # Find best template
        best_template, similarity = self.find_best_template(input_text, category)
        if not best_template:
            print(f"Error: No matching template found for category {category}")
            return None
        
        print(f"  Best match: {best_template} (similarity: {similarity:.2f})")
        
        # Generate report
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = f"{input_file.stem}_comparison_{timestamp}.docx"
        report_path = output_path / report_name
        
        template_text = self.templates[best_template]
        diff_count = self.generate_comparison_report(
            template_text, input_text, str(report_path), input_file.name
        )
        
        print(f"  Report saved: {report_path}")
        
        return AnalysisResult(
            input_file=str(input_file),
            matched_template=best_template,
            similarity_score=similarity,
            category=category,
            differences_count=diff_count,
            report_path=str(report_path),
            timestamp=timestamp
        )
    
    def analyze_batch(self, input_paths: List[str], output_dir: str = "reports", 
                     max_workers: int = 4) -> List[AnalysisResult]:
        """Analyze multiple contracts in parallel"""
        results = []
        
        print(f"\nBatch processing {len(input_paths)} contracts...")
        print(f"Using {max_workers} parallel workers")
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_path = {
                executor.submit(self.analyze_single_contract, path, output_dir): path 
                for path in input_paths
            }
            
            for future in concurrent.futures.as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    result = future.result()
                    if result:
                        results.append(result)
                except Exception as e:
                    print(f"Error processing {path}: {e}")
        
        return results
    
    def generate_batch_summary(self, results: List[AnalysisResult], output_path: str = "batch_summary.json"):
        """Generate a summary report for batch processing"""
        summary = {
            "timestamp": datetime.now().isoformat(),
            "total_processed": len(results),
            "by_category": {},
            "by_template": {},
            "results": []
        }
        
        for result in results:
            # Count by category
            if result.category not in summary["by_category"]:
                summary["by_category"][result.category] = 0
            summary["by_category"][result.category] += 1
            
            # Count by template
            if result.matched_template not in summary["by_template"]:
                summary["by_template"][result.matched_template] = 0
            summary["by_template"][result.matched_template] += 1
            
            # Add result details
            summary["results"].append({
                "file": result.input_file,
                "category": result.category,
                "template": result.matched_template,
                "similarity": result.similarity_score,
                "changes": result.differences_count,
                "report": result.report_path
            })
        
        # Save summary
        with open(output_path, 'w') as f:
            json.dump(summary, f, indent=2)
        
        print(f"\nBatch summary saved to: {output_path}")
        return summary


def main():
    """Main entry point with CLI interface"""
    parser = argparse.ArgumentParser(
        description="Advanced Contract Analyzer with batch processing capabilities"
    )
    parser.add_argument(
        "input", 
        nargs="+", 
        help="Path(s) to contract file(s) or directory containing contracts"
    )
    parser.add_argument(
        "--templates", 
        default="templates",
        help="Directory containing template files (default: templates)"
    )
    parser.add_argument(
        "--output", 
        default="reports",
        help="Output directory for reports (default: reports)"
    )
    parser.add_argument(
        "--batch",
        action="store_true",
        help="Enable batch processing mode"
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=4,
        help="Number of parallel workers for batch processing (default: 4)"
    )
    parser.add_argument(
        "--model",
        default="all-MiniLM-L6-v2",
        help="Sentence transformer model to use (default: all-MiniLM-L6-v2)"
    )
    
    args = parser.parse_args()
    
    # Initialize analyzer
    analyzer = ContractAnalyzer(template_dir=args.templates, model_name=args.model)
    
    # Collect input files
    input_files = []
    for input_path in args.input:
        path = Path(input_path)
        if path.is_file() and path.suffix == ".docx":
            input_files.append(str(path))
        elif path.is_dir():
            # Add all .docx files in directory
            input_files.extend([str(f) for f in path.glob("*.docx")])
    
    if not input_files:
        print("Error: No valid .docx files found")
        return 1
    
    # Process files
    if len(input_files) == 1 and not args.batch:
        # Single file mode
        result = analyzer.analyze_single_contract(input_files[0], args.output)
        if result:
            print(f"\nAnalysis complete!")
            print(f"Changes detected: {result.differences_count}")
    else:
        # Batch mode
        results = analyzer.analyze_batch(input_files, args.output, args.workers)
        if results:
            analyzer.generate_batch_summary(results, 
                                          os.path.join(args.output, "batch_summary.json"))
            print(f"\nBatch analysis complete!")
            print(f"Successfully processed: {len(results)}/{len(input_files)} files")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())