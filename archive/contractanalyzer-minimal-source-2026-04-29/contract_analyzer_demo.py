#!/usr/bin/env python3
"""
Contract Analyzer Demo - Lightweight version without ML dependencies
This version uses simple text matching instead of semantic similarity
"""

import os
import sys
from pathlib import Path
from datetime import datetime
import docx
import difflib
from typing import Dict, List, Tuple, Optional
import json
import argparse


class ContractAnalyzerDemo:
    """Demo contract analyzer using text matching"""
    
    def __init__(self, template_dir: str = 'templates'):
        self.template_dir = Path(template_dir)
        self.templates = {}
        
        # Keyword definitions
        self.vendor_keywords = {
            "capgemini": "VENDOR_CAPGEMINI",
            "blue optima": "VENDOR_BLUEOPTIMA",
            "blueoptima": "VENDOR_BLUEOPTIMA",
            "epam": "VENDOR_EPAM",
        }
        
        self.type_keywords = {
            "statement of work": "TYPE_SOW",
            "sow": "TYPE_SOW",
            "change order": "TYPE_CHANGEORDER",
            "change request": "TYPE_CHANGEORDER",
        }
        
        self.load_templates()
    
    def load_templates(self):
        """Load all templates from the template directory"""
        if not self.template_dir.exists():
            print(f"Warning: Template directory '{self.template_dir}' does not exist. Creating it...")
            self.template_dir.mkdir(parents=True)
            return
        
        template_files = list(self.template_dir.glob("*.docx"))
        if not template_files:
            print(f"Warning: No .docx templates found in '{self.template_dir}'")
            print("\nTo use this tool, you need to:")
            print("1. Generate templates using the provided prompt")
            print("2. Place them in the 'templates' directory")
            print("3. Run the analyzer again")
            return
        
        print(f"Loading {len(template_files)} templates...")
        for template_path in template_files:
            text = self.extract_text_from_docx(str(template_path))
            if text:
                self.templates[template_path.name] = text
                print(f"  - Loaded: {template_path.name}")
        
        print(f"Successfully loaded {len(self.templates)} templates")
    
    def extract_text_from_docx(self, filepath: str, skip_initial_table: bool = True) -> str:
        """Extract text from a .docx file"""
        try:
            doc = docx.Document(filepath)
            
            if not doc.paragraphs and not doc.tables:
                return ""
            
            if not doc.tables or not skip_initial_table:
                return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            
            # Skip initial table
            body_paragraphs = []
            start_collecting = False
            
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    para = docx.text.paragraph.Paragraph(element, doc)
                    if start_collecting and para.text.strip():
                        body_paragraphs.append(para.text.strip())
                elif element.tag.endswith('tbl'):
                    start_collecting = True
            
            return "\n".join(body_paragraphs)
            
        except Exception as e:
            print(f"Error reading {filepath}: {e}")
            return ""
    
    def classify_document(self, text: str) -> Optional[str]:
        """Classify document based on keywords"""
        if not text:
            return None
            
        lower_text = text.lower()
        
        # Check vendor keywords first
        for keyword, category in self.vendor_keywords.items():
            if keyword in lower_text:
                return category
        
        # Check document type keywords
        for keyword, category in self.type_keywords.items():
            if keyword in lower_text:
                return category
        
        return None
    
    def find_best_template(self, input_text: str, category: str) -> Tuple[str, float]:
        """Find best matching template using text similarity"""
        relevant_templates = {
            name: text for name, text in self.templates.items()
            if name.upper().startswith(category)
        }
        
        if not relevant_templates:
            return None, 0.0
        
        # Simple text similarity using difflib
        best_match = None
        best_score = 0.0
        
        for template_name, template_text in relevant_templates.items():
            # Calculate similarity ratio
            matcher = difflib.SequenceMatcher(None, template_text, input_text)
            similarity = matcher.ratio()
            
            if similarity > best_score:
                best_score = similarity
                best_match = template_name
        
        return best_match, best_score
    
    def generate_comparison_report(self, base_text: str, new_text: str, 
                                 output_path: str, contract_name: str) -> int:
        """Generate Word-style track changes comparison report"""
        doc = docx.Document()
        
        # Add header
        header = doc.add_heading('Contract Comparison Report', 0)
        header.alignment = 1  # Center
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Contract: ").bold = True
        metadata_para.add_run(f"{contract_name}\n")
        metadata_para.add_run(f"Generated: ").bold = True
        metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        metadata_para.add_run(f"Template Comparison: ").bold = True
        metadata_para.add_run("Approved Template vs. Submitted Contract\n")
        doc.add_paragraph()
        
        # Calculate differences first
        base_lines = base_text.splitlines()
        new_lines = new_text.splitlines()
        matcher = difflib.SequenceMatcher(None, base_lines, new_lines)
        
        added_count = 0
        deleted_count = 0
        modified_count = 0
        total_changes = 0
        
        # Count changes for summary
        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == 'insert':
                added_count += j2 - j1
                total_changes += j2 - j1
            elif opcode == 'delete':
                deleted_count += i2 - i1
                total_changes += i2 - i1
            elif opcode == 'replace':
                modified_count += max(i2 - i1, j2 - j1)
                total_changes += max(i2 - i1, j2 - j1)
        
        # Add summary section with track changes legend
        doc.add_heading('Change Summary & Legend', level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Changes Detected: {total_changes}\n").bold = True
        
        # Legend with actual formatting examples
        legend_para = doc.add_paragraph()
        legend_para.add_run("Legend:\n").bold = True
        
        # Deleted text example
        deleted_run = legend_para.add_run("■ Deleted Text")
        deleted_run.font.strike = True
        deleted_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
        legend_para.add_run(f" - Text removed from template ({deleted_count} deletions)\n")
        
        # Inserted text example  
        inserted_run = legend_para.add_run("■ Inserted Text")
        inserted_run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)
        inserted_run.font.underline = True
        legend_para.add_run(f" - Text added to contract ({added_count} additions)\n")
        
        # Modified text example
        modified_run = legend_para.add_run("■ Modified Text")
        modified_run.font.color.rgb = docx.shared.RGBColor(255, 165, 0)
        legend_para.add_run(f" - Text changed from template ({modified_count} modifications)\n")
        
        doc.add_paragraph()
        
        # Add detailed comparison with improved formatting
        doc.add_heading('Detailed Track Changes View', level=1)
        
        # Create sentence-level comparison for better readability
        base_sentences = self._split_into_sentences(base_text)
        new_sentences = self._split_into_sentences(new_text)
        sentence_matcher = difflib.SequenceMatcher(None, base_sentences, new_sentences)
        
        comparison_para = doc.add_paragraph()
        
        for opcode, i1, i2, j1, j2 in sentence_matcher.get_opcodes():
            if opcode == 'equal':
                # Unchanged text - normal formatting
                text = ' '.join(base_sentences[i1:i2])
                if text.strip():
                    comparison_para.add_run(text + ' ')
                    
            elif opcode == 'delete':
                # Deleted text - red strikethrough
                text = ' '.join(base_sentences[i1:i2])
                if text.strip():
                    run = comparison_para.add_run(text + ' ')
                    run.font.strike = True
                    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                    
            elif opcode == 'insert':
                # Inserted text - green underline
                text = ' '.join(new_sentences[j1:j2])
                if text.strip():
                    run = comparison_para.add_run(text + ' ')
                    run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)
                    run.font.underline = True
                    
            elif opcode == 'replace':
                # Modified text - show both versions
                base_text_part = ' '.join(base_sentences[i1:i2])
                new_text_part = ' '.join(new_sentences[j1:j2])
                
                if base_text_part.strip():
                    # Show deleted version
                    run = comparison_para.add_run(base_text_part + ' ')
                    run.font.strike = True
                    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                
                if new_text_part.strip():
                    # Show inserted version
                    run = comparison_para.add_run(new_text_part + ' ')
                    run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)
                    run.font.underline = True
        
        # Add section break for better readability
        doc.add_paragraph()
        doc.add_heading('Change Analysis', level=1)
        
        analysis_para = doc.add_paragraph()
        if total_changes == 0:
            run = analysis_para.add_run("✓ COMPLIANT: No unauthorized changes detected.")
            run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)
            run.bold = True
        elif total_changes <= 10:
            run = analysis_para.add_run("⚠ MINOR CHANGES: Few modifications detected - review recommended.")
            run.font.color.rgb = docx.shared.RGBColor(255, 165, 0)
            run.bold = True
        else:
            run = analysis_para.add_run("❌ MAJOR CHANGES: Significant modifications detected - thorough review required.")
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            run.bold = True
        
        # Save document
        doc.save(output_path)
        return total_changes
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better diff granularity"""
        import re
        # Simple sentence splitting - can be enhanced with nltk if needed
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def analyze_contract(self, input_path: str, output_dir: str = "reports"):
        """Analyze a single contract"""
        input_file = Path(input_path)
        if not input_file.exists():
            print(f"Error: File '{input_path}' not found")
            return None
        
        print(f"\nAnalyzing: {input_file.name}")
        
        # Extract text
        input_text = self.extract_text_from_docx(str(input_file))
        if not input_text:
            print(f"Error: Could not extract text from {input_file.name}")
            return None
        
        # Classify
        category = self.classify_document(input_text)
        if not category:
            print(f"Warning: Could not classify {input_file.name}")
            print("The document doesn't match any known vendor or type keywords.")
            return None
        
        print(f"  Category: {category}")
        
        # Find template
        best_template, similarity = self.find_best_template(input_text, category)
        if not best_template:
            print(f"Error: No matching template found for category {category}")
            return None
        
        print(f"  Best match: {best_template} (similarity: {similarity:.2%})")
        
        # Generate report
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = f"{input_file.stem}_comparison_{timestamp}.docx"
        report_path = output_path / report_name
        
        template_text = self.templates[best_template]
        diff_count = self.generate_comparison_report(
            template_text, input_text, str(report_path), input_file.name
        )
        
        print(f"  Report saved: {report_path}")
        print(f"  Changes detected: {diff_count}")
        
        # Add compliance status to console output
        if diff_count == 0:
            print(f"  Status: ✓ COMPLIANT (No unauthorized changes)")
        elif diff_count <= 10:
            print(f"  Status: ⚠ MINOR CHANGES (Review recommended)")
        else:
            print(f"  Status: ❌ MAJOR CHANGES (Thorough review required)")
        
        return {
            "file": str(input_file),
            "category": category,
            "template": best_template,
            "similarity": similarity,
            "changes": diff_count,
            "report": str(report_path)
        }


def main():
    parser = argparse.ArgumentParser(
        description="Contract Analyzer Demo - Compare contracts against templates"
    )
    parser.add_argument(
        "input",
        nargs="*",
        help="Contract file(s) to analyze"
    )
    parser.add_argument(
        "--templates",
        default="templates",
        help="Template directory (default: templates)"
    )
    parser.add_argument(
        "--output",
        default="reports",
        help="Output directory (default: reports)"
    )
    parser.add_argument(
        "--demo",
        action="store_true",
        help="Run demo mode with sample data"
    )
    
    args = parser.parse_args()
    
    # Initialize analyzer
    analyzer = ContractAnalyzerDemo(template_dir=args.templates)
    
    if not analyzer.templates:
        print("\n" + "="*60)
        print("NO TEMPLATES FOUND!")
        print("="*60)
        print("\nTo use this analyzer:")
        print("1. Use the AI prompt in 'generate_templates_and_contracts_prompt.md'")
        print("2. Generate the 5 templates and 20 test contracts")
        print("3. Place templates in the 'templates' directory")
        print("4. Run the analyzer on the test contracts")
        return 1
    
    if args.demo or not args.input:
        print("\n" + "="*60)
        print("DEMO MODE")
        print("="*60)
        print("\nAvailable templates:")
        for template in analyzer.templates:
            print(f"  - {template}")
        
        print("\nTo analyze contracts, run:")
        print(f"  python {sys.argv[0]} <contract_file.docx>")
        print("\nOr for batch processing:")
        print(f"  python {sys.argv[0]} contracts/*.docx")
        return 0
    
    # Process input files
    results = []
    for input_file in args.input:
        if Path(input_file).suffix == ".docx":
            result = analyzer.analyze_contract(input_file, args.output)
            if result:
                results.append(result)
    
    if results:
        print("\n" + "="*60)
        print("ANALYSIS SUMMARY")
        print("="*60)
        print(f"Total contracts analyzed: {len(results)}")
        
        # Group by template
        by_template = {}
        for r in results:
            template = r['template']
            if template not in by_template:
                by_template[template] = []
            by_template[template].append(r)
        
        print("\nBy template:")
        for template, contracts in by_template.items():
            print(f"  {template}: {len(contracts)} contracts")
            total_changes = sum(c['changes'] for c in contracts)
            if total_changes > 0:
                print(f"    WARNING: {total_changes} total changes detected!")
        
        # Save summary
        summary_path = Path(args.output) / "analysis_summary.json"
        with open(summary_path, 'w') as f:
            json.dump({
                "timestamp": datetime.now().isoformat(),
                "total_analyzed": len(results),
                "results": results
            }, f, indent=2)
        print(f"\nDetailed summary saved to: {summary_path}")


if __name__ == "__main__":
    main()