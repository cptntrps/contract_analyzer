#!/usr/bin/env python3

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_analysis_data():
    """Load the analysis results from JSON file"""
    with open('analysis_results.json', 'r') as f:
        return json.load(f)

def generate_comparison_table_report():
    """Generate a Word document with before/after comparison table"""
    print("Generating comparison table report...")
    
    data = load_analysis_data()
    doc = Document()
    
    # Title
    title = doc.add_heading('Contract Analysis - Before vs After Comparison', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Analysis Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"Contract: ").bold = True
    summary.add_run(f"{data[0]['contract']}\n")
    summary.add_run(f"Template: ").bold = True
    summary.add_run(f"{data[0]['template']}\n")
    summary.add_run(f"Total Changes: ").bold = True
    summary.add_run(f"{data[0]['changes']}\n")
    summary.add_run(f"Similarity: ").bold = True
    summary.add_run(f"{data[0]['similarity']}%\n")
    summary.add_run(f"Status: ").bold = True
    summary.add_run(f"{data[0]['status'].upper()}\n")
    
    # Changes breakdown
    doc.add_heading('Changes Breakdown', level=1)
    
    # Count changes by classification
    significant_count = len([c for c in data[0]['analysis'] if c['classification'] == 'SIGNIFICANT'])
    critical_count = len([c for c in data[0]['analysis'] if c['classification'] == 'CRITICAL'])
    inconsequential_count = len([c for c in data[0]['analysis'] if c['classification'] == 'INCONSEQUENTIAL'])
    
    breakdown = doc.add_paragraph()
    breakdown.add_run(f"CRITICAL Changes: ").bold = True
    breakdown.add_run(f"{critical_count}\n")
    breakdown.add_run(f"SIGNIFICANT Changes: ").bold = True
    breakdown.add_run(f"{significant_count}\n")
    breakdown.add_run(f"INCONSEQUENTIAL Changes: ").bold = True
    breakdown.add_run(f"{inconsequential_count}\n")
    
    # Detailed changes table
    doc.add_heading('Detailed Changes', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Classification'
    hdr_cells[1].text = 'Template'
    hdr_cells[2].text = 'Document'
    hdr_cells[3].text = 'Relevance'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add change rows
    for change in data[0]['analysis']:
        row_cells = table.add_row().cells
        row_cells[0].text = change['classification']
        row_cells[1].text = change['deleted_text'][:200] + "..." if len(change['deleted_text']) > 200 else change['deleted_text']
        row_cells[2].text = change['inserted_text'][:200] + "..." if len(change['inserted_text']) > 200 else change['inserted_text']
        row_cells[3].text = change['explanation'][:300] + "..." if len(change['explanation']) > 300 else change['explanation']
    
    # Save the document
    doc.save('Contract_Analysis_Comparison_Table.docx')
    print("✓ Comparison table report saved as: Contract_Analysis_Comparison_Table.docx")

def generate_text_analysis_report():
    """Generate a text-based analysis report"""
    print("Generating text analysis report...")
    
    data = load_analysis_data()
    
    report = f"""
CONTRACT ANALYSIS REPORT
{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

EXECUTIVE SUMMARY
================
Contract: {data[0]['contract']}
Template: {data[0]['template']}
Analysis Date: {data[0]['date']}
Total Changes Identified: {data[0]['changes']}
Similarity Score: {data[0]['similarity']}%
Risk Status: {data[0]['status'].upper()}

CHANGE CLASSIFICATION SUMMARY
============================
"""
    
    # Count changes by classification
    classifications = {}
    for change in data[0]['analysis']:
        classification = change['classification']
        classifications[classification] = classifications.get(classification, 0) + 1
    
    for classification, count in classifications.items():
        report += f"{classification}: {count} changes\n"
    
    report += "\nCRITICAL CHANGES (HIGH RISK)\n"
    report += "=" * 50 + "\n"
    
    critical_changes = [c for c in data[0]['analysis'] if c['classification'] == 'CRITICAL']
    if critical_changes:
        for i, change in enumerate(critical_changes, 1):
            report += f"\n{i}. {change['explanation']}\n"
            report += f"   TEMPLATE: {change['deleted_text'][:100]}...\n"
            report += f"   DOCUMENT: {change['inserted_text'][:100]}...\n"
    else:
        report += "No critical changes identified.\n"
    
    report += "\nSIGNIFICANT CHANGES (MEDIUM RISK)\n"
    report += "=" * 50 + "\n"
    
    significant_changes = [c for c in data[0]['analysis'] if c['classification'] == 'SIGNIFICANT']
    if significant_changes:
        for i, change in enumerate(significant_changes, 1):
            report += f"\n{i}. {change['explanation']}\n"
            report += f"   TEMPLATE: {change['deleted_text'][:100]}...\n"
            report += f"   DOCUMENT: {change['inserted_text'][:100]}...\n"
    else:
        report += "No significant changes identified.\n"
    
    report += "\nRECOMMENDATIONS\n"
    report += "=" * 50 + "\n"
    
    if data[0]['status'] == 'critical':
        report += "⚠️  CRITICAL: This contract requires immediate legal review before execution.\n"
        report += "   - Multiple high-risk changes detected\n"
        report += "   - Recommend full legal and compliance review\n"
        report += "   - Do not proceed without approval\n"
    elif len(significant_changes) > 0:
        report += "⚠️  SIGNIFICANT: This contract has material changes that need review.\n"
        report += "   - Review all significant changes with legal team\n"
        report += "   - Verify financial and scope implications\n"
        report += "   - Obtain necessary approvals before signing\n"
    else:
        report += "✅ LOW RISK: Changes appear to be routine updates.\n"
        report += "   - Standard review process can be followed\n"
        report += "   - No major red flags identified\n"
    
    # Save the report
    with open('Contract_Analysis_Text_Report.txt', 'w') as f:
        f.write(report)
    
    print("✓ Text analysis report saved as: Contract_Analysis_Text_Report.txt")

def generate_template_comparison_list():
    """Generate a list comparing template vs contract"""
    print("Generating template vs contract comparison list...")
    
    data = load_analysis_data()
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Template vs Contract Comparison', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"Template: ").bold = True
    summary.add_run(f"{data[0]['template']}\n")
    summary.add_run(f"Contract: ").bold = True
    summary.add_run(f"{data[0]['contract']}\n")
    summary.add_run(f"Match Rate: ").bold = True
    summary.add_run(f"{data[0]['similarity']}%\n")
    
    # Changes by category
    doc.add_heading('Changes by Category', level=1)
    
    categories = {
        'Financial Terms': [],
        'Scope of Work': [],
        'Dates and Deadlines': [],
        'Personnel': [],
        'Legal/Compliance': [],
        'Other': []
    }
    
    # Categorize changes
    for change in data[0]['analysis']:
        explanation = change['explanation'].lower()
        if any(keyword in explanation for keyword in ['value', 'cost', 'price', 'payment', 'financial', '$']):
            categories['Financial Terms'].append(change)
        elif any(keyword in explanation for keyword in ['scope', 'service', 'deliverable', 'objective', 'work']):
            categories['Scope of Work'].append(change)
        elif any(keyword in explanation for keyword in ['date', 'deadline', 'timeline', 'commence', 'conclude']):
            categories['Dates and Deadlines'].append(change)
        elif any(keyword in explanation for keyword in ['name', 'person', 'prepared', 'approved', 'client']):
            categories['Personnel'].append(change)
        elif any(keyword in explanation for keyword in ['legal', 'compliance', 'terms', 'conditions']):
            categories['Legal/Compliance'].append(change)
        else:
            categories['Other'].append(change)
    
    # Add each category
    for category, changes in categories.items():
        if changes:
            doc.add_heading(f'{category} ({len(changes)} changes)', level=2)
            for i, change in enumerate(changes, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(f"{change['explanation']}\n")
                para.add_run("   Template: ").bold = True
                para.add_run(f"{change['deleted_text'][:150]}...\n")
                para.add_run("   Document: ").bold = True
                para.add_run(f"{change['inserted_text'][:150]}...\n")
    
    # Save the document
    doc.save('Template_vs_Contract_Comparison.docx')
    print("✓ Template vs contract comparison saved as: Template_vs_Contract_Comparison.docx")

def main():
    """Generate all reports"""
    print("🔥 GENERATING USEFUL CONTRACT ANALYSIS REPORTS 🔥")
    print("=" * 60)
    
    try:
        generate_comparison_table_report()
        generate_text_analysis_report()
        generate_template_comparison_list()
        
        print("\n" + "=" * 60)
        print("✅ ALL REPORTS GENERATED SUCCESSFULLY!")
        print("=" * 60)
        print("📄 Files created:")
        print("   1. Contract_Analysis_Comparison_Table.docx - Detailed before/after table")
        print("   2. Contract_Analysis_Text_Report.txt - Executive summary and analysis")
        print("   3. Template_vs_Contract_Comparison.docx - Categorized comparison list")
        print("   4. reports/ directory - Original Word redlined documents with track changes")
        print("\n🌐 Dashboard: http://localhost:5000")
        print("=" * 60)
        
    except Exception as e:
        print(f"Error generating reports: {e}")
        print("Make sure analysis_results.json exists and is valid.")

if __name__ == "__main__":
    main() 