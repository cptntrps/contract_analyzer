#!/usr/bin/env python3
"""
Enhanced Report Generator for Contract Analysis
Generates three types of documents:
1. Redlined Document (.docx) - Word format with track changes
2. Changes Table (.xlsx and optionally .docx) - Structured table
3. Summary Report (.pdf and optionally .docx) - Plain language summary
"""

import json
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

# Document generation imports
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, Color
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

# Windows-only COM interface for Word
try:
    import win32com.client as win32
    import pythoncom
    COM_AVAILABLE = True
except ImportError:
    COM_AVAILABLE = False

import logging
logger = logging.getLogger(__name__)

class EnhancedReportGenerator:
    """Enhanced report generator with multiple output formats"""
    
    def __init__(self, reports_dir: str = "reports"):
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(exist_ok=True)
        
        # Initialize generated reports tracking
        self.generated_reports = []
        
        # Check if Windows COM is available
        self.com_available = COM_AVAILABLE
        if not self.com_available:
            logger.warning("Windows COM interface not available. Word track changes feature disabled.")
    
    def generate_word_com_redlined_document(self, analysis_data: Dict[str, Any], template_path: str, base_name: str) -> Optional[str]:
        """Generate Word document with proper track changes using COM interface (Windows only)"""
        
        if not self.com_available:
            logger.warning("Windows COM interface not available. Cannot generate Word track changes document.")
            return None
            
        # COM objects to clean up
        word = None
        original = None
        revised = None
        compared_doc = None
        
        try:
            # Initialize COM
            pythoncom.CoInitialize()
            
            # Get original contract path from analysis data
            contract_path = analysis_data.get('contract_path')
            if not contract_path:
                raise ValueError("Contract path not found in analysis data. Please re-run the analysis.")
            
            # Output path for redlined document
            output_path = self.reports_dir / f"{base_name}_word_com_redlined.docx"
            
            # Verify files exist before proceeding
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Template file not found: {template_path}")
            if not Path(contract_path).exists():
                raise FileNotFoundError(f"Contract file not found: {contract_path}")
            
            logger.info(f"Starting Word COM comparison: {Path(template_path).name} vs {Path(contract_path).name}")
            
            # Use Word COM to compare documents
            word = win32.Dispatch('Word.Application')
            word.Visible = False
            
            # Open both documents
            logger.debug("Opening original document...")
            original = word.Documents.Open(str(Path(template_path).absolute()))
            logger.debug("Opening revised document...")
            revised = word.Documents.Open(str(Path(contract_path).absolute()))
            
            # Perform the comparison with track changes
            logger.debug("Performing document comparison...")
            
            # Try different parameter combinations for better compatibility
            try:
                # First, try the full parameter set
                compared_doc = word.CompareDocuments(
                    OriginalDocument=original,
                    RevisedDocument=revised,
                    Destination=1,  # wdCompareDestinationNew
                    Granularity=1,  # wdGranularityWordLevel
                    CompareFormatting=True,
                    CompareCaseChanges=True,
                    CompareWhitespace=True,
                    CompareTables=True,
                    CompareHeaders=True,
                    CompareFootnotes=True,
                    CompareTextboxes=True,
                    CompareFields=True,
                    CompareComments=True
                )
            except Exception as e:
                logger.debug(f"Full parameter set failed: {e}. Trying simplified version...")
                # If that fails, try with only essential parameters
                try:
                    compared_doc = word.CompareDocuments(
                        OriginalDocument=original,
                        RevisedDocument=revised,
                        Destination=1,
                        Granularity=1,
                        CompareFormatting=True
                    )
                except Exception as e2:
                    logger.debug(f"Simplified version failed: {e2}. Trying minimal parameters...")
                    # If that fails, try minimal parameters
                    compared_doc = word.CompareDocuments(
                        OriginalDocument=original,
                        RevisedDocument=revised,
                        Destination=1
                    )
            
            # Verify the comparison succeeded
            if compared_doc is None:
                raise Exception("Document comparison failed - returned None")
            
            logger.debug("Configuring track changes display...")
            # Configure track changes display for proper color scheme
            compared_doc.TrackRevisions = True
            compared_doc.ShowRevisions = True
            
            # Set review pane display options for proper Word color scheme
            view = compared_doc.ActiveWindow.View
            view.RevisionsView = 0  # wdRevisionsViewFinal - show final with markup
            view.ShowRevisionsAndComments = True
            
            # Configure revision display options using Word's comparison colors
            try:
                view.ShowInsertionsAndDeletions = True
                view.ShowFormatChanges = True
                view.ShowComments = True
                
                # Set Word's comparison colors using different approach
                revision_options = compared_doc.Application.Options
                
                # Use Word color constants (different from previous approach)
                # wdColorRed = 6, wdColorBlue = 5, wdColorGreen = 4, wdColorAutomatic = 0
                revision_options.InsertedTextColor = 6    # wdColorRed for insertions
                revision_options.DeletedTextColor = 5     # wdColorBlue for deletions  
                revision_options.RevisedLinesColor = 6    # wdColorRed for revised lines
                
                # Set proper marking styles
                revision_options.InsertedTextMark = 2     # wdInsertedTextMarkUnderline
                revision_options.DeletedTextMark = 1      # wdDeletedTextMarkStrikeThrough
                
                logger.debug("Applied Word comparison colors: Red insertions, Blue deletions")
            except Exception as color_error:
                logger.debug(f"Could not set revision colors (will use Word defaults): {color_error}")
                # Continue without custom colors - Word will use its defaults
            
            # Save the redlined version
            logger.debug(f"Saving redlined document to: {output_path}")
            compared_doc.SaveAs2(str(output_path.absolute()))
            
            logger.info(f"Word COM redlined document generated: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error generating Word COM redlined document: {e}")
            return None
            
        finally:
            # Safe COM cleanup - don't propagate cleanup errors
            self._safe_com_cleanup(word, original, revised, compared_doc)
            
            # Always cleanup COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass
    
    def _safe_com_cleanup(self, word=None, original=None, revised=None, compared_doc=None):
        """Safely clean up COM objects without propagating errors"""
        
        # Close documents in reverse order (newest first)
        documents_to_close = [
            ("compared_doc", compared_doc),
            ("revised", revised), 
            ("original", original)
        ]
        
        for doc_name, doc in documents_to_close:
            if doc is not None:
                try:
                    logger.debug(f"Closing {doc_name}...")
                    doc.Close(False)  # Don't save changes
                except Exception as e:
                    logger.debug(f"Warning: Could not close {doc_name}: {e}")
                    # Continue cleanup even if individual document close fails
        
        # Quit Word application
        if word is not None:
            try:
                logger.debug("Quitting Word application...")
                word.Quit()
            except Exception as e:
                logger.debug(f"Warning: Could not quit Word cleanly: {e}")
                # This is expected for COM disconnection errors
                # Don't treat as a failure since documents were already saved
        
        logger.debug("COM cleanup completed")
    
    def _create_contract_document(self, analysis_data: Dict[str, Any], output_path: Path):
        """Create a contract document from analysis data for comparison"""
        
        doc = Document()
        
        # Add contract title
        title = doc.add_heading(f"Contract: {analysis_data['contract']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contract content based on analysis
        contract_para = doc.add_paragraph()
        contract_para.add_run("This contract represents the analyzed version with all changes applied.\n\n")
        
        # Add changes organized by criticality
        changes_by_type = self._group_changes_by_type(analysis_data.get('analysis', []))
        
        for criticality in ['critical', 'significant', 'inconsequential']:
            if criticality in changes_by_type and changes_by_type[criticality]:
                # Add section heading
                doc.add_heading(f"{criticality.title()} Changes", level=1)
                
                for change in changes_by_type[criticality]:
                    # Add the "after" text (how it appears in the contract)
                    change_para = doc.add_paragraph()
                    change_para.add_run(f"• {change.get('after', change.get('description', 'No description'))}")
        
        # Save the document
        doc.save(output_path)

    def generate_all_reports(self, analysis_data: Dict[str, Any], 
                           report_types: Dict[str, bool] = None,
                           template_path: str = None) -> Dict[str, List[str]]:
        """
        Generate all requested report types
        
        Args:
            analysis_data: Analysis result data
            report_types: Dict specifying which reports to generate
            template_path: Path to template file for Word COM comparison
            
        Returns:
            Dict mapping report types to file paths
        """
        if report_types is None:
            report_types = {
                'review_docx': True,
                'changes_xlsx': True,
                'changes_docx': False,
                'summary_pdf': True,
                'summary_docx': False,
                'word_com_redlined': True
            }
        
        generated_files = {
            'review_documents': [],
            'changes_tables': [],
            'summary_reports': [],
            'word_com_redlined': []
        }
        
        base_name = f"{analysis_data['id']}"
        
        try:
            # 1. Generate Redlined Document (.docx)
            if report_types.get('review_docx', True):
                review_path = self.generate_review_document(analysis_data, base_name)
                generated_files['review_documents'].append(review_path)
                
            # 2. Generate Changes Table (.xlsx and optionally .docx)
            if report_types.get('changes_xlsx', True):
                changes_xlsx_path = self.generate_changes_table_xlsx(analysis_data, base_name)
                generated_files['changes_tables'].append(changes_xlsx_path)
                
            if report_types.get('changes_docx', False):
                changes_docx_path = self.generate_changes_table_docx(analysis_data, base_name)
                generated_files['changes_tables'].append(changes_docx_path)
                
            # 3. Summary report generation removed
                
            # 4. Generate Word COM Redlined Document (Windows only)
            if report_types.get('word_com_redlined', True) and template_path:
                word_com_path = self.generate_word_com_redlined_document(analysis_data, template_path, base_name)
                if word_com_path:
                    generated_files['word_com_redlined'].append(word_com_path)
                
            # Update metadata
            self.generated_reports.append({
                'id': analysis_data['id'],
                'contract': analysis_data['contract'],
                'timestamp': datetime.now().isoformat(),
                'files': generated_files
            })
            
            return generated_files
            
        except Exception as e:
            print(f"Error generating reports: {e}")
            return {}
    
    def generate_review_document(self, analysis_data: Dict[str, Any], base_name: str) -> str:
        """Generate Redlined Document (.docx) with Word-like track changes format"""
        
        doc = Document()
        
        # Document title
        title = doc.add_heading('Contract Redlined Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive summary section
        doc.add_heading('Executive Summary', level=1)
        
        # Summary content in paragraph format (not table)
        summary_para = doc.add_paragraph()
        summary_para.add_run("Contract: ").bold = True
        summary_para.add_run(f"{analysis_data['contract']}\n")
        summary_para.add_run("Template Used: ").bold = True
        summary_para.add_run(f"{analysis_data['template']}\n")
        summary_para.add_run("Analysis Date: ").bold = True
        summary_para.add_run(f"{analysis_data['date']}\n")
        summary_para.add_run("Total Changes Detected: ").bold = True
        summary_para.add_run(f"{analysis_data['changes']}\n")
        summary_para.add_run("Similarity Score: ").bold = True
        summary_para.add_run(f"{analysis_data['similarity']}%\n")
        summary_para.add_run("Review Status: ").bold = True
        summary_para.add_run(f"{analysis_data['status']}")
        
        # Add space
        doc.add_paragraph()
        
        # Risk assessment
        doc.add_heading('Risk Assessment', level=1)
        risk_level = self._determine_risk_level(analysis_data)
        risk_para = doc.add_paragraph()
        risk_para.add_run("Overall Risk Level: ").bold = True
        risk_run = risk_para.add_run(risk_level)
        risk_run.font.bold = True
        
        if risk_level == "HIGH":
            risk_run.font.color.rgb = RGBColor(220, 53, 69)
        elif risk_level == "MEDIUM":
            risk_run.font.color.rgb = RGBColor(255, 193, 7)
        else:
            risk_run.font.color.rgb = RGBColor(25, 135, 84)
        
        # Add risk explanation
        doc.add_paragraph()
        risk_explanation = doc.add_paragraph()
        risk_explanation.add_run(self._get_risk_explanation(risk_level, analysis_data))
        
        # Changes organized by criticality (highest to lowest)
        doc.add_heading('Changes by Criticality', level=1)
        
        # Group changes by type
        changes_by_type = self._group_changes_by_type(analysis_data.get('analysis', []))
        
        # Process in order of criticality: Critical -> Significant -> Inconsequential
        criticality_order = [
            ('critical', 'CRITICAL', RGBColor(220, 53, 69)),
            ('significant', 'SIGNIFICANT', RGBColor(255, 193, 7)),
            ('inconsequential', 'INCONSEQUENTIAL', RGBColor(25, 135, 84))
        ]
        
        for criticality_key, criticality_label, criticality_color in criticality_order:
            changes = changes_by_type.get(criticality_key, [])
            if changes:
                # Section header
                section_heading = doc.add_heading(f'{criticality_label} Changes ({len(changes)})', level=2)
                section_heading.runs[0].font.color.rgb = criticality_color
                
                # Add section description
                if criticality_key == 'critical':
                    desc = doc.add_paragraph()
                    desc.add_run("⚠️ These changes may significantly impact contract terms and require immediate review.")
                    desc.runs[0].font.color.rgb = criticality_color
                elif criticality_key == 'significant':
                    desc = doc.add_paragraph()
                    desc.add_run("⚠️ These changes may affect contract performance and should be reviewed.")
                    desc.runs[0].font.color.rgb = criticality_color
                else:
                    desc = doc.add_paragraph()
                    desc.add_run("✓ These changes are minor and typically routine.")
                    desc.runs[0].font.color.rgb = criticality_color
                
                doc.add_paragraph()  # Space after description
                
                # List all changes in this criticality level
                for i, change in enumerate(changes, 1):
                    # Change number and explanation
                    change_para = doc.add_paragraph()
                    change_para.add_run(f"{i}. ").bold = True
                    change_para.add_run(change.get('explanation', 'Change detected'))
                    
                    # Create a more Word-like redline format
                    redline_para = doc.add_paragraph()
                    redline_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # Show deleted text (strikethrough effect with red)
                    deleted_text = change.get('deleted_text', '').strip()
                    if deleted_text:
                        redline_para.add_run("REMOVED: ").bold = True
                        deleted_run = redline_para.add_run(deleted_text)
                        deleted_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                        deleted_run.font.strike = True
                        redline_para.add_run("\n")
                    
                    # Show inserted text (green)
                    inserted_text = change.get('inserted_text', '').strip()
                    if inserted_text:
                        redline_para.add_run("ADDED: ").bold = True
                        inserted_run = redline_para.add_run(inserted_text)
                        inserted_run.font.color.rgb = RGBColor(25, 135, 84)  # Green
                        inserted_run.font.underline = True
                    
                    # Add spacing between changes
                    doc.add_paragraph()
        
        # Summary and recommendations
        doc.add_heading('Summary and Recommendations', level=1)
        
        # Count changes by type for summary
        total_critical = len(changes_by_type.get('critical', []))
        total_significant = len(changes_by_type.get('significant', []))
        total_inconsequential = len(changes_by_type.get('inconsequential', []))
        
        summary_text = doc.add_paragraph()
        summary_text.add_run("Change Summary:\n").bold = True
        summary_text.add_run(f"• Critical changes: {total_critical}\n")
        summary_text.add_run(f"• Significant changes: {total_significant}\n")
        summary_text.add_run(f"• Inconsequential changes: {total_inconsequential}\n")
        
        doc.add_paragraph()
        
        # Recommendations
        recommendations = self._generate_recommendations(analysis_data)
        rec_heading = doc.add_paragraph()
        rec_heading.add_run("Recommendations:\n").bold = True
        
        for i, rec in enumerate(recommendations, 1):
            rec_para = doc.add_paragraph()
            rec_para.add_run(f"{i}. ").bold = True
            rec_para.add_run(rec)
        
        # Save document
        file_path = self.reports_dir / f"{base_name}_redlined_document.docx"
        doc.save(file_path)
        
        return str(file_path)
    
    def generate_changes_table_xlsx(self, analysis_data: Dict[str, Any], base_name: str) -> str:
        """Generate Changes Table (.xlsx) with structured comparison"""
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Changes Analysis"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        # Border style
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Headers
        headers = [
            "Change #", "Section", "Change Type", "Classification", 
            "Template", "Document", "Relevance", 
            "Risk Level", "Recommendation"
        ]
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # Data rows
        changes = analysis_data.get('analysis', [])
        for row, change in enumerate(changes, 2):
            row_data = [
                row - 1,  # Change number
                self._extract_section(change.get('deleted_text', '')),
                self._classify_change_type(change),
                change.get('classification', 'UNKNOWN'),
                change.get('deleted_text', '')[:100] + '...' if len(change.get('deleted_text', '')) > 100 else change.get('deleted_text', ''),
                change.get('inserted_text', '')[:100] + '...' if len(change.get('inserted_text', '')) > 100 else change.get('inserted_text', ''),
                change.get('explanation', 'No explanation provided'),
                self._get_change_risk_level(change),
                self._get_change_recommendation(change)
            ]
            
            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                
                # Color coding based on classification
                if col == 4:  # Classification column
                    classification = change.get('classification', '').upper()
                    if classification == 'CRITICAL':
                        cell.fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
                    elif classification == 'SIGNIFICANT':
                        cell.fill = PatternFill(start_color="FFF4E6", end_color="FFF4E6", fill_type="solid")
                    else:
                        cell.fill = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid")
        
        # Adjust column widths
        column_widths = [8, 15, 15, 15, 25, 25, 30, 12, 25]
        for col, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width
        
        # Add summary sheet
        summary_ws = wb.create_sheet("Summary")
        summary_ws.append(["Contract Analysis Summary"])
        summary_ws.append([])
        summary_ws.append(["Contract", analysis_data['contract']])
        summary_ws.append(["Template", analysis_data['template']])
        summary_ws.append(["Analysis Date", analysis_data['date']])
        summary_ws.append(["Total Changes", analysis_data['changes']])
        summary_ws.append(["Similarity Score", f"{analysis_data['similarity']}%"])
        summary_ws.append(["Status", analysis_data['status']])
        
        # Save workbook
        file_path = self.reports_dir / f"{base_name}_changes_table.xlsx"
        wb.save(file_path)
        
        return str(file_path)
    
    def generate_changes_table_docx(self, analysis_data: Dict[str, Any], base_name: str) -> str:
        """Generate Changes Table (.docx) version"""
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Contract Changes Table', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Contract: ").bold = True
        summary_para.add_run(f"{analysis_data['contract']}\n")
        summary_para.add_run(f"Template: ").bold = True
        summary_para.add_run(f"{analysis_data['template']}\n")
        summary_para.add_run(f"Changes: ").bold = True
        summary_para.add_run(f"{analysis_data['changes']}\n")
        summary_para.add_run(f"Similarity: ").bold = True
        summary_para.add_run(f"{analysis_data['similarity']}%\n")
        
        # Changes table
        doc.add_heading('Detailed Changes', level=1)
        
        changes = analysis_data.get('analysis', [])
        if changes:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header
            headers = ["#", "Classification", "Template", "Document", "Relevance"]
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for i, change in enumerate(changes, 1):
                row = table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = change.get('classification', 'UNKNOWN')
                row.cells[2].text = change.get('deleted_text', '')[:100] + '...' if len(change.get('deleted_text', '')) > 100 else change.get('deleted_text', '')
                row.cells[3].text = change.get('inserted_text', '')[:100] + '...' if len(change.get('inserted_text', '')) > 100 else change.get('inserted_text', '')
                row.cells[4].text = change.get('explanation', 'No explanation provided')
        
        # Save document
        file_path = self.reports_dir / f"{base_name}_changes_table.docx"
        doc.save(file_path)
        
        return str(file_path)
    
    # Summary report generation methods removed
    
    # Helper methods
    def _group_changes_by_type(self, changes: List[Dict]) -> Dict[str, List[Dict]]:
        """Group changes by type"""
        groups = {
            'critical': [],
            'significant': [],
            'inconsequential': []
        }
        
        for change in changes:
            classification = change.get('classification', 'INCONSEQUENTIAL').lower()
            if classification in groups:
                groups[classification].append(change)
            else:
                groups['inconsequential'].append(change)
        
        return groups
    
    def _determine_risk_level(self, analysis_data: Dict[str, Any]) -> str:
        """
        Determine overall risk level using minimum criticality rule
        
        New Logic:
        - Any critical change = HIGH risk (minimum of critical)
        - Placeholder → Actual content = Generally low impact
        - Actual → Different actual content = Critical (high risk)
        """
        changes = analysis_data.get('analysis', [])
        critical_count = len([c for c in changes if c.get('classification') == 'CRITICAL'])
        significant_count = len([c for c in changes if c.get('classification') == 'SIGNIFICANT'])
        
        # Minimum criticality rule: One critical change makes entire contract critical
        if critical_count > 0:
            return "HIGH"
        elif significant_count > 5:
            return "HIGH"
        elif significant_count > 0:
            return "MEDIUM"
        else:
            return "LOW"
    
    def _get_risk_explanation(self, risk_level: str, analysis_data: Dict[str, Any]) -> str:
        """Get explanation for risk level based on new criticality logic"""
        changes = analysis_data.get('analysis', [])
        critical_count = len([c for c in changes if c.get('classification') == 'CRITICAL'])
        
        if risk_level == "HIGH":
            if critical_count > 0:
                return f"This contract contains {critical_count} critical change(s) involving actual value-to-value modifications (e.g., price changes, service changes, company changes). These require immediate legal review and approval before proceeding."
            else:
                return "This contract contains numerous significant changes that require immediate legal review and approval before proceeding."
        elif risk_level == "MEDIUM":
            return "This contract contains significant changes that should be reviewed by legal counsel before execution."
        else:
            return "This contract has minimal changes (mostly placeholder → actual content) and can proceed through standard review processes."
    
    def _generate_recommendations(self, analysis_data: Dict[str, Any]) -> List[str]:
        """Generate recommendations based on analysis and new criticality logic"""
        recommendations = []
        
        changes = analysis_data.get('analysis', [])
        critical_changes = [c for c in changes if c.get('classification') == 'CRITICAL']
        risk_level = self._determine_risk_level(analysis_data)
        
        if risk_level == "HIGH":
            recommendations.extend([
                "🚨 CRITICAL: Schedule immediate legal review with qualified counsel",
                "🚨 Do not execute contract until all critical changes are approved",
                "🚨 Focus on value-to-value changes (price, service, company modifications)"
            ])
            
            # Add specific recommendations based on critical change types
            if critical_changes:
                monetary_changes = [c for c in critical_changes if any(symbol in str(c.get('deleted_text', '')) + str(c.get('inserted_text', '')) for symbol in ['$', '€', '£', '¥', 'USD', 'EUR'])]
                if monetary_changes:
                    recommendations.append("💰 Verify all monetary value changes and get financial approval")
                
                service_changes = [c for c in critical_changes if any(keyword in str(c.get('deleted_text', '')).lower() + str(c.get('inserted_text', '')).lower() for keyword in ['service', 'product', 'deliverable', 'work'])]
                if service_changes:
                    recommendations.append("🔄 Review all service/product changes for scope alignment")
                
                company_changes = [c for c in critical_changes if len([w for w in str(c.get('deleted_text', '')).split() if w and w[0].isupper()]) >= 2]
                if company_changes:
                    recommendations.append("🏢 Verify all company/entity changes for accuracy")
        
        elif risk_level == "MEDIUM":
            recommendations.extend([
                "⚠️ Review significant changes with legal team",
                "⚠️ Verify all financial and timeline implications",
                "⚠️ Obtain necessary internal approvals before signing"
            ])
        else:
            recommendations.extend([
                "✅ Proceed with standard contract review process",
                "✅ Most changes are placeholder → actual content (low risk)",
                "✅ Verify routine changes align with business requirements",
                "✅ Consider expedited approval for minimal risk contracts"
            ])
        
        return recommendations
    
    def _extract_section(self, text: str) -> str:
        """Extract section from text"""
        # Simple section extraction - can be enhanced
        if "Section" in text:
            return text.split("Section")[1].split()[0] if "Section" in text else "Unknown"
        return "General"
    
    def _classify_change_type(self, change: Dict[str, Any]) -> str:
        """Classify the type of change"""
        explanation = change.get('explanation', '').lower()
        
        if any(word in explanation for word in ['financial', 'cost', 'price', 'payment', 'value']):
            return "Financial"
        elif any(word in explanation for word in ['date', 'timeline', 'deadline', 'schedule']):
            return "Timeline"
        elif any(word in explanation for word in ['scope', 'deliverable', 'service']):
            return "Scope"
        elif any(word in explanation for word in ['party', 'name', 'entity']):
            return "Parties"
        else:
            return "Other"
    
    def _get_change_risk_level(self, change: Dict[str, Any]) -> str:
        """Get risk level for individual change"""
        classification = change.get('classification', 'INCONSEQUENTIAL').upper()
        
        if classification == 'CRITICAL':
            return "HIGH"
        elif classification == 'SIGNIFICANT':
            return "MEDIUM"
        else:
            return "LOW"
    
    def _get_change_recommendation(self, change: Dict[str, Any]) -> str:
        """Get recommendation for individual change"""
        risk_level = self._get_change_risk_level(change)
        
        if risk_level == "HIGH":
            return "Requires legal approval"
        elif risk_level == "MEDIUM":
            return "Review with stakeholders"
        else:
            return "Standard approval process"

# Example usage
def main():
    """Example usage of the enhanced report generator"""
    
    # Load sample analysis data
    sample_data = {
        'id': 'Contract_001_Sample_20240115',
        'contract': 'Sample Contract.docx',
        'template': 'Standard Template.docx',
        'status': 'Changes - Moderate',
        'changes': 25,
        'similarity': 78.5,
        'date': '2024-01-15T10:30:00',
        'analysis': [
            {
                'classification': 'SIGNIFICANT',
                'explanation': 'Financial value changed from $100,000 to $150,000',
                'deleted_text': 'Total contract value: $100,000',
                'inserted_text': 'Total contract value: $150,000'
            },
            {
                'classification': 'INCONSEQUENTIAL',
                'explanation': 'Minor formatting change',
                'deleted_text': 'This Agreement',
                'inserted_text': 'This Contract'
            }
        ]
    }
    
    # Generate reports
    generator = EnhancedReportGenerator()
    
    report_types = {
        'review_docx': True,
        'changes_xlsx': True,
        'changes_docx': True,
        'summary_pdf': True,
        'summary_docx': True
    }
    
    generated_files = generator.generate_all_reports(sample_data, report_types)
    
    print("Generated Reports:")
    for report_type, files in generated_files.items():
        print(f"\n{report_type.title()}:")
        for file_path in files:
            print(f"  - {file_path}")

if __name__ == "__main__":
    main() 