"""
Unit tests for DocumentProcessor service
"""

import pytest
from unittest.mock import Mock, patch, mock_open
from pathlib import Path
import tempfile
import os

from app.core.services.document_processor import DocumentProcessor


class TestDocumentProcessor:
    """Test suite for DocumentProcessor service"""

    def test_init(self):
        """Test processor initialization"""
        processor = DocumentProcessor()
        assert processor is not None

    def test_extract_text_from_docx_success(self, test_docx_file):
        """Test successful text extraction from DOCX file"""
        processor = DocumentProcessor()
        text = processor.extract_text_from_docx(str(test_docx_file))
        
        assert isinstance(text, str)
        assert len(text) > 0
        assert 'Test Contract' in text

    def test_extract_text_from_docx_file_not_found(self):
        """Test text extraction with non-existent file"""
        processor = DocumentProcessor()
        text = processor.extract_text_from_docx('nonexistent.docx')
        assert text == ""

    def test_extract_text_from_docx_invalid_file(self, tmp_path):
        """Test text extraction with invalid DOCX file"""
        processor = DocumentProcessor()
        
        # Create a text file with .docx extension
        invalid_file = tmp_path / "invalid.docx"
        invalid_file.write_text("This is not a valid DOCX file")
        
        text = processor.extract_text_from_docx(str(invalid_file))
        assert text == ""

    def test_extract_text_from_docx_empty_document(self):
        """Test text extraction from empty document"""
        processor = DocumentProcessor()
        
        # Create an empty DOCX document
        from docx import Document
        doc = Document()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            try:
                text = processor.extract_text_from_docx(tmp.name)
                assert isinstance(text, str)
                assert len(text.strip()) == 0
            finally:
                if os.path.exists(tmp.name):
                    os.unlink(tmp.name)

    def test_clean_text(self):
        """Test text cleaning functionality"""
        processor = DocumentProcessor()
        
        # Test with various whitespace and formatting issues
        dirty_text = "  This   has\n\n\nextra   whitespace\t\tand\r\nline breaks  "
        clean_text = processor.clean_text(dirty_text)
        
        assert clean_text == "This has extra whitespace and line breaks"

    def test_clean_text_empty(self):
        """Test cleaning empty text"""
        processor = DocumentProcessor()
        
        assert processor.clean_text("") == ""
        assert processor.clean_text("   ") == ""
        assert processor.clean_text("\n\n\n") == ""

    def test_clean_text_unicode(self):
        """Test cleaning text with unicode characters"""
        processor = DocumentProcessor()
        
        unicode_text = "  This has üñïçödé characters  "
        clean_text = processor.clean_text(unicode_text)
        
        assert clean_text == "This has üñïçödé characters"

    def test_normalize_text(self):
        """Test text normalization"""
        processor = DocumentProcessor()
        
        # Test case normalization
        mixed_case = "ThIs Is MiXeD cAsE tExT"
        normalized = processor.normalize_text(mixed_case)
        
        assert normalized == "this is mixed case text"

    def test_normalize_text_punctuation(self):
        """Test normalization with punctuation"""
        processor = DocumentProcessor()
        
        text_with_punct = "Hello, World! How are you?"
        normalized = processor.normalize_text(text_with_punct)
        
        # Should remove punctuation and convert to lowercase
        assert "," not in normalized
        assert "!" not in normalized
        assert "?" not in normalized
        assert normalized == "hello world how are you"

    def test_validate_file_path_valid(self, test_docx_file):
        """Test file path validation with valid file"""
        processor = DocumentProcessor()
        
        assert processor.validate_file_path(str(test_docx_file)) == True

    def test_validate_file_path_nonexistent(self):
        """Test file path validation with non-existent file"""
        processor = DocumentProcessor()
        
        assert processor.validate_file_path("/nonexistent/file.docx") == False

    def test_validate_file_path_wrong_extension(self, tmp_path):
        """Test file path validation with wrong extension"""
        processor = DocumentProcessor()
        
        # Create a file with wrong extension
        wrong_ext_file = tmp_path / "test.txt"
        wrong_ext_file.write_text("test content")
        
        assert processor.validate_file_path(str(wrong_ext_file)) == False

    def test_validate_file_path_empty(self):
        """Test file path validation with empty path"""
        processor = DocumentProcessor()
        
        assert processor.validate_file_path("") == False
        assert processor.validate_file_path(None) == False

    def test_get_file_info(self, test_docx_file):
        """Test getting file information"""
        processor = DocumentProcessor()
        
        info = processor.get_file_info(str(test_docx_file))
        
        assert isinstance(info, dict)
        assert 'size' in info
        assert 'extension' in info
        assert 'filename' in info
        assert info['extension'] == '.docx'
        assert info['size'] > 0

    def test_get_file_info_nonexistent(self):
        """Test getting file info for non-existent file"""
        processor = DocumentProcessor()
        
        info = processor.get_file_info("/nonexistent/file.docx")
        
        assert info is None

    def test_process_document_success(self, test_docx_file):
        """Test complete document processing"""
        processor = DocumentProcessor()
        
        result = processor.process_document(str(test_docx_file))
        
        assert isinstance(result, dict)
        assert 'text' in result
        assert 'cleaned_text' in result
        assert 'normalized_text' in result
        assert 'file_info' in result
        
        assert len(result['text']) > 0
        assert isinstance(result['file_info'], dict)

    def test_process_document_error(self):
        """Test document processing with error"""
        processor = DocumentProcessor()
        
        result = processor.process_document("/nonexistent/file.docx")
        
        assert result is None

    @patch('app.core.services.document_processor.Document')
    def test_extract_text_exception_handling(self, mock_document):
        """Test exception handling in text extraction"""
        processor = DocumentProcessor()
        
        # Mock Document to raise an exception
        mock_document.side_effect = Exception("Test exception")
        
        text = processor.extract_text_from_docx("test.docx")
        assert text == ""

    def test_extract_paragraphs(self, test_docx_file):
        """Test extracting paragraphs separately"""
        processor = DocumentProcessor()
        
        paragraphs = processor.extract_paragraphs(str(test_docx_file))
        
        assert isinstance(paragraphs, list)
        assert len(paragraphs) > 0
        assert all(isinstance(p, str) for p in paragraphs)

    def test_extract_paragraphs_error(self):
        """Test paragraph extraction with error"""
        processor = DocumentProcessor()
        
        paragraphs = processor.extract_paragraphs("/nonexistent/file.docx")
        assert paragraphs == []

    def test_get_document_structure(self, test_docx_file):
        """Test getting document structure information"""
        processor = DocumentProcessor()
        
        structure = processor.get_document_structure(str(test_docx_file))
        
        assert isinstance(structure, dict)
        assert 'paragraph_count' in structure
        assert 'word_count' in structure
        assert 'character_count' in structure
        assert structure['paragraph_count'] > 0

    def test_get_document_structure_error(self):
        """Test document structure extraction with error"""
        processor = DocumentProcessor()
        
        structure = processor.get_document_structure("/nonexistent/file.docx")
        assert structure is None


class TestDocumentProcessorEdgeCases:
    """Test edge cases and error conditions"""

    def test_very_large_document(self):
        """Test processing very large document"""
        processor = DocumentProcessor()
        
        # Create a large document
        from docx import Document
        doc = Document()
        
        # Add many paragraphs
        for i in range(1000):
            doc.add_paragraph(f"This is paragraph {i} with some content to make it longer.")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            try:
                text = processor.extract_text_from_docx(tmp.name)
                assert isinstance(text, str)
                assert len(text) > 0
                assert "paragraph 500" in text
            finally:
                if os.path.exists(tmp.name):
                    os.unlink(tmp.name)

    def test_document_with_tables(self):
        """Test processing document with tables"""
        processor = DocumentProcessor()
        
        # Create document with table
        from docx import Document
        doc = Document()
        doc.add_paragraph("Document with table")
        
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            try:
                text = processor.extract_text_from_docx(tmp.name)
                assert isinstance(text, str)
                assert "Document with table" in text
                # Tables should be included in text extraction
                assert "Cell 1" in text
            finally:
                if os.path.exists(tmp.name):
                    os.unlink(tmp.name)

    def test_document_with_special_formatting(self):
        """Test processing document with special formatting"""
        processor = DocumentProcessor()
        
        # Create document with various formatting
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('Document Title', 0)
        doc.add_heading('Section Header', level=1)
        
        p = doc.add_paragraph('This paragraph has ')
        p.add_run('bold').bold = True
        p.add_run(' and ')
        p.add_run('italic').italic = True
        p.add_run(' text.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            
            try:
                text = processor.extract_text_from_docx(tmp.name)
                assert isinstance(text, str)
                assert "Document Title" in text
                assert "Section Header" in text
                assert "bold" in text
                assert "italic" in text
            finally:
                if os.path.exists(tmp.name):
                    os.unlink(tmp.name)