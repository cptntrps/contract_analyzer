"""
Pytest configuration and fixtures for Contract Analyzer tests
"""

import pytest
import tempfile
import shutil
import os
from pathlib import Path
from unittest.mock import Mock, patch
import json

# Import test utilities
from tests import TEST_DATA_DIR, TEST_UPLOAD_DIR, TEST_REPORTS_DIR, TEST_TEMPLATES_DIR

# Import application modules (updated for new architecture)
from app.config.settings import get_config
from app.core.services.analyzer import ContractAnalyzer
from app.services.llm.providers.openai import OpenAIProvider
from app.services.reports.generator import ReportGenerator
from app.utils.security.validators import SecurityValidator
from app.config.user_settings import UserSettingsManager
from app import create_app

@pytest.fixture(scope="session")
def test_config():
    """Test configuration with safe defaults"""
    return {
        'UPLOAD_FOLDER': str(TEST_UPLOAD_DIR),
        'REPORTS_FOLDER': str(TEST_REPORTS_DIR),
        'TEMPLATES_FOLDER': str(TEST_TEMPLATES_DIR),
        'MAX_CONTENT_LENGTH': 16 * 1024 * 1024,  # 16MB
        'ALLOWED_EXTENSIONS': ['docx'],
        'DEBUG': True,
        'TESTING': True,
        'SECRET_KEY': 'test-secret-key',
        'FLASK_HOST': '127.0.0.1',
        'FLASK_PORT': 5001,  # Different from production
        'LOG_LEVEL': 'DEBUG',
        'OPENAI_API_KEY': 'test-api-key',
        'OPENAI_MODEL': 'gpt-4o',
        'LLM_PROVIDER': 'openai',
        'LLM_TEMPERATURE': 0.1,
        'LLM_MAX_TOKENS': 2000,
        'LLM_TIMEOUT': 30
    }

@pytest.fixture
def mock_config(test_config):
    """Mock configuration for testing"""
    # Create a mock config object with the test config attributes
    mock_config_obj = Mock()
    for key, value in test_config.items():
        setattr(mock_config_obj, key.upper(), value)
    
    # Patch the new configuration system
    with patch('app.config.settings.get_config') as mock_get_config:
        mock_get_config.return_value = mock_config_obj
        yield mock_config_obj

@pytest.fixture
def sample_contract_text():
    """Sample contract text for testing"""
    return """
    CONSULTING AGREEMENT
    
    This Consulting Agreement ("Agreement") is entered into on January 1, 2025,
    between Acme Corporation ("Company") and John Smith ("Consultant").
    
    1. SCOPE OF WORK
    The Consultant agrees to provide software development services as described
    in the attached Statement of Work.
    
    2. COMPENSATION
    The Company will pay the Consultant $150 per hour for services rendered.
    Payment will be made within 30 days of invoice receipt.
    
    3. TERMINATION
    Either party may terminate this agreement with 30 days written notice.
    
    4. CONFIDENTIALITY
    The Consultant agrees to maintain confidentiality of all proprietary information.
    """

@pytest.fixture
def sample_template_text():
    """Sample template text for testing"""
    return """
    CONSULTING AGREEMENT
    
    This Consulting Agreement ("Agreement") is entered into on [DATE],
    between [COMPANY NAME] ("Company") and [CONSULTANT NAME] ("Consultant").
    
    1. SCOPE OF WORK
    The Consultant agrees to provide [SERVICE TYPE] services as described
    in the attached Statement of Work.
    
    2. COMPENSATION
    The Company will pay the Consultant $[HOURLY_RATE] per hour for services rendered.
    Payment will be made within [PAYMENT_TERMS] days of invoice receipt.
    
    3. TERMINATION
    Either party may terminate this agreement with [NOTICE_PERIOD] days written notice.
    
    4. CONFIDENTIALITY
    The Consultant agrees to maintain confidentiality of all proprietary information.
    """

@pytest.fixture
def sample_changes():
    """Sample changes for testing analysis"""
    return [
        ('delete', '[DATE]'),
        ('insert', 'January 1, 2025'),
        ('delete', '[COMPANY NAME]'),
        ('insert', 'Acme Corporation'),
        ('delete', '[CONSULTANT NAME]'),
        ('insert', 'John Smith'),
        ('delete', '[SERVICE TYPE]'),
        ('insert', 'software development'),
        ('delete', '$[HOURLY_RATE]'),
        ('insert', '$150'),
        ('delete', '[PAYMENT_TERMS]'),
        ('insert', '30'),
        ('delete', '[NOTICE_PERIOD]'),
        ('insert', '30')
    ]

@pytest.fixture
def sample_analysis_result():
    """Sample analysis result with multi-stakeholder data"""
    return {
        'explanation': 'Updated hourly rate from placeholder to $150',
        'category': 'FINANCIAL',
        'classification': 'SIGNIFICANT',
        'financial_impact': 'DIRECT',
        'required_reviews': ['FINANCE_APPROVAL', 'LEGAL_REVIEW'],
        'procurement_flags': ['rate_change', 'vendor_performance_impact'],
        'review_priority': 'high',
        'deleted_text': '$[HOURLY_RATE]',
        'inserted_text': '$150',
        'confidence': 'high'
    }

@pytest.fixture
def mock_llm_provider():
    """Mock LLM provider for testing"""
    from app.services.llm.providers.base import BaseLLMProvider, LLMResponse
    
    provider = Mock(spec=BaseLLMProvider)
    
    # Mock the response
    mock_response = LLMResponse(
        content='{"explanation": "Test analysis result", "classification": "SIGNIFICANT"}',
        usage={'total_tokens': 100},
        model='test-model',
        provider='test',
        response_time=1.0
    )
    provider.generate_response.return_value = mock_response
    provider.is_healthy.return_value = True
    provider.get_available_models.return_value = ['test-model', 'test-model-2']
    provider.model = 'test-model'
    
    return provider


@pytest.fixture
def mock_openai_provider():
    """Mock OpenAI provider for testing"""
    provider = Mock(spec=OpenAIProvider)
    provider.check_connection.return_value = True
    provider.get_current_model.return_value = 'gpt-4o'
    provider.get_available_models.return_value = [
        {'name': 'gpt-4o', 'description': 'Test GPT-4o model', 'current': True, 'recommended': True}
    ]
    provider._generate_response.return_value = '{"explanation": "Test response"}'
    return provider

@pytest.fixture
def test_docx_file():
    """Create a test DOCX file"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Contract', 0)
    doc.add_paragraph('This is a test contract for testing purposes.')
    doc.add_paragraph('It contains sample text that can be analyzed.')
    
    test_file = TEST_DATA_DIR / 'test_contract.docx'
    test_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(test_file))
    
    yield test_file
    
    # Cleanup
    if test_file.exists():
        test_file.unlink()

@pytest.fixture
def test_template_file():
    """Create a test template DOCX file"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Template', 0)
    doc.add_paragraph('This is a test template for testing purposes.')
    doc.add_paragraph('It contains [PLACEHOLDER] text that should be replaced.')
    
    test_file = TEST_DATA_DIR / 'test_template.docx'
    test_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(test_file))
    
    yield test_file
    
    # Cleanup
    if test_file.exists():
        test_file.unlink()

@pytest.fixture
def analyzer():
    """Contract analyzer instance"""
    config = {
        'llm_settings': {
            'provider': 'openai',
            'api_key': 'test-key',
            'model': 'gpt-4o',
            'temperature': 0.1
        }
    }
    return ContractAnalyzer(config)

@pytest.fixture
def security_validator():
    """Security validator instance"""
    return SecurityValidator()

@pytest.fixture
def user_config_manager():
    """User config manager instance"""
    return UserSettingsManager()

@pytest.fixture
def report_generator():
    """Report generator instance"""
    config = {'REPORTS_FOLDER': str(TEST_REPORTS_DIR)}
    return ReportGenerator(reports_dir=str(TEST_REPORTS_DIR), config=config)

@pytest.fixture
def flask_app(mock_config):
    """Flask application instance for testing"""
    with patch('app.config.settings.get_config') as mock_get_config:
        mock_get_config.return_value = mock_config
        app = create_app('testing')
        app.config['TESTING'] = True
        yield app

@pytest.fixture
def client(flask_app):
    """Flask test client"""
    return flask_app.test_client()

@pytest.fixture(autouse=True)
def cleanup_after_test():
    """Cleanup after each test"""
    yield
    # Clean up any test files that might have been created
    for test_dir in [TEST_UPLOAD_DIR, TEST_REPORTS_DIR, TEST_TEMPLATES_DIR]:
        if test_dir.exists():
            for file in test_dir.glob('*'):
                if file.is_file():
                    file.unlink()

@pytest.fixture
def mock_file_upload():
    """Mock file upload data"""
    return {
        'filename': 'test_contract.docx',
        'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'size': 1024
    }

@pytest.fixture
def sample_contract_data():
    """Sample contract data for testing"""
    return {
        'id': 'test-contract-123',
        'name': 'Test Contract.docx',
        'path': str(TEST_DATA_DIR / 'test_contract.docx'),
        'type': 'Generic',
        'size': 1024,
        'uploaded': '2025-01-01T00:00:00Z'
    }

@pytest.fixture
def sample_template_data():
    """Sample template data for testing"""
    return {
        'id': 'test-template-123',
        'name': 'Test Template.docx',
        'path': str(TEST_DATA_DIR / 'test_template.docx'),
        'category': 'Generic',
        'size': 1024,
        'uploaded': '2025-01-01T00:00:00Z'
    }

@pytest.fixture
def sample_analysis_data():
    """Sample analysis data for testing"""
    return {
        'id': 'test-analysis-123',
        'contract': 'Test Contract.docx',
        'contract_path': str(TEST_DATA_DIR / 'test_contract.docx'),
        'template': 'Test Template.docx',
        'template_path': str(TEST_DATA_DIR / 'test_template.docx'),
        'status': 'Changes - Minor',
        'changes': 5,
        'similarity': 85.5,
        'date': '2025-01-01T00:00:00Z',
        'analysis': [
            {
                'explanation': 'Filled placeholder with actual value',
                'classification': 'INCONSEQUENTIAL',
                'category': 'ADMINISTRATIVE',
                'financial_impact': 'NONE',
                'required_reviews': ['ROUTINE'],
                'procurement_flags': [],
                'review_priority': 'normal',
                'confidence': 'high'
            }
        ]
    }

# Test data generators
def generate_test_contract_text(placeholders=None):
    """Generate test contract text with optional placeholders"""
    if placeholders is None:
        placeholders = {
            'company_name': 'Test Company Inc.',
            'date': 'January 1, 2025',
            'amount': '$1,000'
        }
    
    return f"""
    CONSULTING AGREEMENT
    
    This agreement is made on {placeholders['date']} between {placeholders['company_name']}
    and the service provider.
    
    The total compensation will be {placeholders['amount']} for the services provided.
    """

def generate_test_template_text():
    """Generate test template text with placeholders"""
    return """
    CONSULTING AGREEMENT
    
    This agreement is made on [DATE] between [COMPANY_NAME]
    and the service provider.
    
    The total compensation will be [AMOUNT] for the services provided.
    """