"""
Unit tests for report formatters
"""

import pytest
import tempfile
import os
import json
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

from app.services.reports.formatters.excel import ExcelFormatter
from app.services.reports.formatters.word import WordFormatter
from app.services.reports.formatters.pdf import PDFFormatter


class TestExcelFormatter:
    """Test suite for ExcelFormatter"""

    def test_init(self):
        """Test formatter initialization"""
        formatter = ExcelFormatter()
        assert formatter is not None

    def test_generate_changes_table_success(self):
        """Test successful Excel changes table generation"""
        formatter = ExcelFormatter()
        
        analysis_data = {
            'contract_name': 'Test Contract.docx',
            'template_name': 'Test Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 85.5,
            'total_changes': 3,
            'analysis_results': [
                {
                    'explanation': 'Date placeholder filled',
                    'classification': 'INCONSEQUENTIAL',
                    'category': 'ADMINISTRATIVE',
                    'deleted_text': '[DATE]',
                    'inserted_text': '2025-01-01',
                    'confidence': 'high'
                },
                {
                    'explanation': 'Company name updated',
                    'classification': 'SIGNIFICANT', 
                    'category': 'BUSINESS',
                    'deleted_text': '[COMPANY]',
                    'inserted_text': 'Acme Corp',
                    'confidence': 'high'
                },
                {
                    'explanation': 'Payment terms modified',
                    'classification': 'CRITICAL',
                    'category': 'FINANCIAL',
                    'deleted_text': '30 days',
                    'inserted_text': '45 days',
                    'confidence': 'medium'
                }
            ]
        }
        
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_changes_table(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            
            # Verify Excel file can be opened
            import openpyxl
            workbook = openpyxl.load_workbook(output_path)
            assert 'Analysis Summary' in workbook.sheetnames
            assert 'Changes' in workbook.sheetnames
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_changes_table_empty_data(self):
        """Test Excel generation with empty analysis data"""
        formatter = ExcelFormatter()
        
        analysis_data = {
            'contract_name': 'Test Contract.docx',
            'template_name': 'Test Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 100.0,
            'total_changes': 0,
            'analysis_results': []
        }
        
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_changes_table(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            
            # Verify Excel file structure
            import openpyxl
            workbook = openpyxl.load_workbook(output_path)
            changes_sheet = workbook['Changes']
            
            # Should have headers but no data rows
            assert changes_sheet['A1'].value == 'Change #'
            assert changes_sheet['A3'].value is None  # No data row
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_changes_table_invalid_path(self):
        """Test Excel generation with invalid output path"""
        formatter = ExcelFormatter()
        
        analysis_data = {'analysis_results': []}
        invalid_path = '/invalid/path/that/does/not/exist.xlsx'
        
        result = formatter.generate_changes_table(analysis_data, invalid_path)
        assert result is None

    def test_format_cell_styling(self):
        """Test Excel cell styling functionality"""
        formatter = ExcelFormatter()
        
        # Test styling methods exist and work
        assert hasattr(formatter, '_create_header_style')
        assert hasattr(formatter, '_create_data_style')
        assert hasattr(formatter, '_get_classification_style')
        
        # Test classification-specific styling
        critical_style = formatter._get_classification_style('CRITICAL')
        significant_style = formatter._get_classification_style('SIGNIFICANT')
        inconsequential_style = formatter._get_classification_style('INCONSEQUENTIAL')
        
        assert critical_style is not None
        assert significant_style is not None
        assert inconsequential_style is not None

    @patch('openpyxl.Workbook')
    def test_generate_changes_table_exception_handling(self, mock_workbook):
        """Test exception handling in Excel generation"""
        formatter = ExcelFormatter()
        
        # Mock workbook to raise exception
        mock_workbook.side_effect = Exception("Test exception")
        
        analysis_data = {'analysis_results': []}
        
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result = formatter.generate_changes_table(analysis_data, output_path)
            assert result is None
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)


class TestWordFormatter:
    """Test suite for WordFormatter"""

    def test_init(self):
        """Test formatter initialization"""
        formatter = WordFormatter()
        assert formatter is not None

    def test_generate_summary_report_success(self):
        """Test successful Word summary report generation"""
        formatter = WordFormatter()
        
        analysis_data = {
            'contract_name': 'Test Contract.docx',
            'template_name': 'Test Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 85.5,
            'total_changes': 2,
            'analysis_results': [
                {
                    'explanation': 'Date updated',
                    'classification': 'SIGNIFICANT',
                    'category': 'ADMINISTRATIVE',
                    'deleted_text': '[DATE]',
                    'inserted_text': '2025-01-01'
                },
                {
                    'explanation': 'Critical payment change',
                    'classification': 'CRITICAL',
                    'category': 'FINANCIAL',
                    'deleted_text': '30 days',
                    'inserted_text': '60 days'
                }
            ]
        }
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_summary_report(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            
            # Verify Word document can be opened
            from docx import Document
            doc = Document(output_path)
            
            # Check document contains expected content
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'Contract Analysis Report' in doc_text
            assert 'Test Contract.docx' in doc_text
            assert 'CRITICAL' in doc_text
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_redlined_document_success(self, test_docx_file):
        """Test successful redlined document generation"""
        formatter = WordFormatter()
        
        analysis_results = [
            {
                'explanation': 'Text modified',
                'classification': 'SIGNIFICANT',
                'deleted_text': 'original text',
                'inserted_text': 'modified text',
                'position': 100
            }
        ]
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            # Note: This test may fail on non-Windows systems without COM
            result_path = formatter.generate_redlined_document(
                str(test_docx_file), 
                analysis_results, 
                output_path
            )
            
            # On non-Windows, should return None or handle gracefully
            if result_path:
                assert result_path == output_path
                assert os.path.exists(output_path)
                
        except Exception as e:
            # Expected on non-Windows systems
            assert 'COM' in str(e) or 'Windows' in str(e)
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_summary_report_empty_data(self):
        """Test Word report generation with empty data"""
        formatter = WordFormatter()
        
        analysis_data = {
            'contract_name': 'Empty Contract.docx',
            'template_name': 'Empty Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 100.0,
            'total_changes': 0,
            'analysis_results': []
        }
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_summary_report(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            
            # Verify document contains appropriate "no changes" message
            from docx import Document
            doc = Document(output_path)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            assert 'no changes' in doc_text.lower() or '100%' in doc_text
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @patch('docx.Document')
    def test_generate_summary_report_exception_handling(self, mock_document):
        """Test exception handling in Word generation"""
        formatter = WordFormatter()
        
        # Mock Document to raise exception
        mock_document.side_effect = Exception("Test exception")
        
        analysis_data = {'analysis_results': []}
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result = formatter.generate_summary_report(analysis_data, output_path)
            assert result is None
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_com_availability_check(self):
        """Test COM availability checking for Windows features"""
        formatter = WordFormatter()
        
        # Test COM availability check
        com_available = formatter._is_com_available()
        
        # Should return boolean
        assert isinstance(com_available, bool)
        
        # On Windows with Office, should be True; otherwise False
        # We can't assert a specific value since it depends on the environment


class TestPDFFormatter:
    """Test suite for PDFFormatter"""

    def test_init(self):
        """Test formatter initialization"""
        formatter = PDFFormatter()
        assert formatter is not None

    def test_generate_summary_report_success(self):
        """Test successful PDF summary report generation"""
        formatter = PDFFormatter()
        
        analysis_data = {
            'contract_name': 'Test Contract.docx',
            'template_name': 'Test Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 87.2,
            'total_changes': 3,
            'analysis_results': [
                {
                    'explanation': 'Date field updated',
                    'classification': 'INCONSEQUENTIAL',
                    'category': 'ADMINISTRATIVE',
                    'deleted_text': '[DATE]',
                    'inserted_text': '2025-01-01'
                },
                {
                    'explanation': 'Payment terms changed',
                    'classification': 'CRITICAL',
                    'category': 'FINANCIAL',
                    'deleted_text': '30 days',
                    'inserted_text': '45 days'
                }
            ]
        }
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_summary_report(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            
            # Basic PDF validation - should start with PDF header
            with open(output_path, 'rb') as f:
                header = f.read(4)
                assert header == b'%PDF'
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_detailed_report_success(self):
        """Test successful PDF detailed report generation"""
        formatter = PDFFormatter()
        
        analysis_data = {
            'contract_name': 'Detailed Contract.docx',
            'template_name': 'Detailed Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 75.0,
            'total_changes': 5,
            'analysis_results': [
                {
                    'explanation': 'Critical change requiring attention',
                    'classification': 'CRITICAL',
                    'category': 'FINANCIAL',
                    'deleted_text': 'Original amount $1000',
                    'inserted_text': 'New amount $2000',
                    'confidence': 'high',
                    'risk_impact': 'High financial impact'
                }
            ]
        }
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_detailed_report(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            
            # Verify PDF header
            with open(output_path, 'rb') as f:
                header = f.read(4)
                assert header == b'%PDF'
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_generate_summary_report_empty_data(self):
        """Test PDF generation with empty analysis data"""
        formatter = PDFFormatter()
        
        analysis_data = {
            'contract_name': 'Empty Contract.docx',
            'template_name': 'Empty Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 100.0,
            'total_changes': 0,
            'analysis_results': []
        }
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result_path = formatter.generate_summary_report(analysis_data, output_path)
            
            assert result_path == output_path
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @patch('reportlab.pdfgen.canvas.Canvas')
    def test_generate_summary_report_exception_handling(self, mock_canvas):
        """Test exception handling in PDF generation"""
        formatter = PDFFormatter()
        
        # Mock Canvas to raise exception
        mock_canvas.side_effect = Exception("Test exception")
        
        analysis_data = {'analysis_results': []}
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result = formatter.generate_summary_report(analysis_data, output_path)
            assert result is None
            
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_styling_and_formatting(self):
        """Test PDF styling and formatting methods"""
        formatter = PDFFormatter()
        
        # Test styling methods exist
        assert hasattr(formatter, '_add_header')
        assert hasattr(formatter, '_add_footer')
        assert hasattr(formatter, '_add_summary_section')
        assert hasattr(formatter, '_add_changes_section')
        
        # Test color mapping for classifications
        critical_color = formatter._get_classification_color('CRITICAL')
        significant_color = formatter._get_classification_color('SIGNIFICANT')
        inconsequential_color = formatter._get_classification_color('INCONSEQUENTIAL')
        
        assert critical_color is not None
        assert significant_color is not None
        assert inconsequential_color is not None
        assert critical_color != significant_color


class TestReportFormattersIntegration:
    """Test integration between different formatters"""

    def test_all_formatters_with_same_data(self):
        """Test all formatters can handle the same analysis data"""
        analysis_data = {
            'contract_name': 'Integration Test Contract.docx',
            'template_name': 'Integration Test Template.docx',
            'analysis_date': '2025-01-01T00:00:00Z',
            'similarity_score': 82.5,
            'total_changes': 4,
            'analysis_results': [
                {
                    'explanation': 'Date placeholder filled',
                    'classification': 'INCONSEQUENTIAL',
                    'category': 'ADMINISTRATIVE',
                    'deleted_text': '[DATE]',
                    'inserted_text': '2025-01-01'
                },
                {
                    'explanation': 'Payment terms updated',
                    'classification': 'CRITICAL',
                    'category': 'FINANCIAL',
                    'deleted_text': '30 days',
                    'inserted_text': '60 days'
                }
            ]
        }
        
        formatters = [
            (ExcelFormatter(), '.xlsx'),
            (WordFormatter(), '.docx'),
            (PDFFormatter(), '.pdf')
        ]
        
        for formatter, extension in formatters:
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp:
                output_path = tmp.name
            
            try:
                if extension == '.xlsx':
                    result = formatter.generate_changes_table(analysis_data, output_path)
                else:
                    result = formatter.generate_summary_report(analysis_data, output_path)
                
                assert result == output_path
                assert os.path.exists(output_path)
                assert os.path.getsize(output_path) > 0
                
            finally:
                if os.path.exists(output_path):
                    os.unlink(output_path)